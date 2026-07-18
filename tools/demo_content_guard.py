"""Show the content guard passing and then catching a text edit.

Run 1 formats a small synthetic paper the normal way.
Run 2 patches the styling step so it also edits one sentence, the way a
buggy formatter might. That run has to abort without writing the DOCX.

Usage: python tools/demo_content_guard.py
"""

from __future__ import annotations

import sys
from pathlib import Path
from tempfile import TemporaryDirectory

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
sys.dont_write_bytecode = True

from docx import Document  # noqa: E402

import paper_format_agent.pipeline as pipeline  # noqa: E402
from paper_format_agent.rules import extract_rules_from_text  # noqa: E402


def make_synthetic_paper(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是摘要内容。")
    doc.add_paragraph("关键词：测试；格式化")
    doc.add_paragraph("一、绪论")
    doc.add_paragraph("这是正文第一段。")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "列2"
    doc.save(path)


def guard_log(result) -> dict:
    for entry in result.logs:
        if entry.get("action") == "content_guard":
            return entry
    return {}


def main() -> int:
    rules = extract_rules_from_text("正文宋体小四，1.25倍行距。")

    with TemporaryDirectory() as td:
        src = Path(td) / "paper.docx"
        make_synthetic_paper(src)

        print("run 1: normal formatting")
        out1 = Path(td) / "formatted.docx"
        result = pipeline.run_pipeline(src, out1, rules, enforce_content_guard=True)
        log = guard_log(result)
        print("  fingerprint before:", log.get("fingerprint_before"))
        print("  fingerprint after: ", log.get("fingerprint_after"))
        print("  content_changed:", result.content_changed)
        print("  formatted DOCX written:", out1.exists())

        print()
        print("run 2: styling step also edits one sentence (injected bug)")
        real_styles = pipeline.apply_final_styles_from_markers

        def styles_that_also_edit_text(doc, rules):
            count = real_styles(doc, rules)
            for p in doc.paragraphs:
                if "正文第一段" in p.text:
                    p.text = p.text.replace("第一段", "第1段")
                    break
            return count

        pipeline.apply_final_styles_from_markers = styles_that_also_edit_text
        out2 = Path(td) / "formatted_should_not_exist.docx"
        guard_error = None
        try:
            pipeline.run_pipeline(src, out2, rules, enforce_content_guard=True)
        except ValueError as e:
            guard_error = str(e)
        finally:
            pipeline.apply_final_styles_from_markers = real_styles
        print("  aborted:", guard_error)
        print("  formatted DOCX written:", out2.exists())

        if guard_error != "content guard failed: non-whitespace content changed":
            print("expected the content guard error, got:", guard_error)
            return 1
        if out2.exists():
            print("guard did not stop the write, this is a bug")
            return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
