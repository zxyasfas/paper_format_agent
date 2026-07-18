"""Check which authored strings survive a formatting run.

Builds four small synthetic DOCX files, runs the pipeline on each with the
content guard disabled so that losses can be diagnosed instead of aborting
the run, and then looks for every authored string in the output document.

Strict success is an exact match. A guard-equivalent match means the string
only survives after the same whitespace and bullet normalization the content
guard itself uses, so it is reported as a warning rather than a pass, and it
is not an independent check. The guard's own changed flag is printed per
document.

Exit code is 1 if any string is lost outright.

Usage: python tools/benchmark_text_survival.py
"""

from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path
from tempfile import TemporaryDirectory

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
sys.dont_write_bytecode = True

from docx import Document  # noqa: E402

import paper_format_agent.pipeline as pipeline  # noqa: E402
from paper_format_agent.quality import normalize_for_content_guard  # noqa: E402
from paper_format_agent.rules import extract_rules_from_text  # noqa: E402

NOT_COVERED = [
    "real Word lists (w:numPr numbering)",
    "field-based tables of contents",
    "merged or nested tables",
    "footnotes",
    "equations (OMML)",
    "text boxes",
    "tracked changes",
    "embedded objects",
]


def build_zh_thesis(path: Path) -> dict:
    authored = {
        "front matter": ["摘要", "本文研究合成文档的排版行为。", "关键词：排版；指纹；测试"],
        "headings": ["一、绪论", "二、相关工作", "2.1 现状分析"],
        "body": [
            "这是绪论的第一段正文，用来检查正文段落是否原样保留。",
            "相关工作部分的正文段落，句子里带标点、数字 123 和 English words。",
        ],
        "captions": ["图1-1 系统结构示意", "表2-1 实验参数对照"],
        "table cells": ["参数", "取值", "阈值", "0.85"],
        "references": ["参考文献", "[1] 王某某. 合成排版研究[J]. 测试学报, 2024."],
    }
    doc = Document()
    for t in authored["front matter"]:
        doc.add_paragraph(t)
    doc.add_paragraph(authored["headings"][0])
    doc.add_paragraph(authored["body"][0])
    doc.add_paragraph(authored["captions"][0])
    doc.add_paragraph(authored["headings"][1])
    doc.add_paragraph(authored["headings"][2])
    doc.add_paragraph(authored["body"][1])
    doc.add_paragraph(authored["captions"][1])
    table = doc.add_table(rows=2, cols=2)
    cells = authored["table cells"]
    table.cell(0, 0).text = cells[0]
    table.cell(0, 1).text = cells[1]
    table.cell(1, 0).text = cells[2]
    table.cell(1, 1).text = cells[3]
    for t in authored["references"]:
        doc.add_paragraph(t)
    doc.save(path)
    return authored


def build_en_paper(path: Path) -> dict:
    authored = {
        "front matter": ["Abstract", "A synthetic paper for the survival benchmark.", "Keywords: formatting; fingerprint"],
        "headings": ["1. Introduction", "2. Method"],
        "body": [
            "This paragraph checks that plain English body text survives.",
            "Numbers 42 and symbols +/- should come through unchanged.",
        ],
        "captions": ["Table 1. Synthetic settings."],
        "table cells": ["name", "value", "alpha", "0.5"],
        "references": ["References", "Smith, J. (2024). Synthetic layouts. Journal of Fake Results, 12(3), 45-67."],
    }
    doc = Document()
    for t in authored["front matter"]:
        doc.add_paragraph(t)
    doc.add_paragraph(authored["headings"][0])
    doc.add_paragraph(authored["body"][0])
    doc.add_paragraph(authored["headings"][1])
    doc.add_paragraph(authored["body"][1])
    doc.add_paragraph(authored["captions"][0])
    table = doc.add_table(rows=2, cols=2)
    cells = authored["table cells"]
    table.cell(0, 0).text = cells[0]
    table.cell(0, 1).text = cells[1]
    table.cell(1, 0).text = cells[2]
    table.cell(1, 1).text = cells[3]
    for t in authored["references"]:
        doc.add_paragraph(t)
    doc.save(path)
    return authored


def build_toc_doc(path: Path) -> dict:
    authored = {
        "manual toc": ["目录", "一、绪论\t1", "二、方法\t2"],
        "headings": ["一、绪论", "二、方法"],
        "body": ["目录文档的正文段落。", "方法部分的正文段落。"],
    }
    doc = Document()
    doc.add_paragraph(authored["headings"][0])
    doc.add_paragraph(authored["body"][0])
    for t in authored["manual toc"]:
        doc.add_paragraph(t)
    doc.add_paragraph(authored["headings"][1])
    doc.add_paragraph(authored["body"][1])
    doc.save(path)
    return authored


def build_bullet_doc(path: Path) -> dict:
    # plain paragraphs with literal bullet characters, not real Word lists
    authored = {
        "bullet-prefixed plain paragraphs": ["• 第一个要点", "• 第二个要点", "- dash item"],
        "body": ["列表文档的普通正文段落。"],
    }
    doc = Document()
    doc.add_paragraph(authored["body"][0])
    for t in authored["bullet-prefixed plain paragraphs"]:
        doc.add_paragraph(t)
    doc.save(path)
    return authored


def output_texts(doc: Document) -> list:
    # the fixtures only use plain unmerged tables; merged cells would be
    # returned once per grid position here and are on the not-covered list
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return [t for t in texts if t.strip()]


def tally(authored: dict, texts: list) -> dict:
    """Match authored strings against output texts.

    Exact matches consume output instances first. Only the leftovers form the
    normalization pool, so one output line can never satisfy two authored
    strings.
    """
    exact_pool = Counter(texts)
    results = {}
    leftovers = {}

    for category, strings in authored.items():
        exact = 0
        remaining = []
        for s in strings:
            if exact_pool[s] > 0:
                exact_pool[s] -= 1
                exact += 1
            else:
                remaining.append(s)
        results[category] = {"exact": exact, "total": len(strings)}
        leftovers[category] = remaining

    norm_pool = Counter()
    for t, count in exact_pool.items():
        if count > 0:
            nt = normalize_for_content_guard(t)
            if nt:
                norm_pool[nt] += count

    for category, remaining in leftovers.items():
        guard_eq = 0
        lost = []
        for s in remaining:
            ns = normalize_for_content_guard(s)
            if ns and norm_pool[ns] > 0:
                norm_pool[ns] -= 1
                guard_eq += 1
            else:
                lost.append(s)
        results[category]["guard_equivalent"] = guard_eq
        results[category]["lost"] = lost

    return results


def main() -> int:
    builders = [
        ("zh_thesis", build_zh_thesis),
        ("en_paper", build_en_paper),
        ("toc_doc", build_toc_doc),
        ("bullet_doc", build_bullet_doc),
    ]
    rules = extract_rules_from_text("正文宋体小四，1.25倍行距。")
    lost_total = 0
    category_rows = 0
    string_total = 0

    with TemporaryDirectory() as td:
        for name, build in builders:
            src = Path(td) / (name + ".docx")
            out = Path(td) / (name + "_formatted.docx")
            authored = build(src)
            result = pipeline.run_pipeline(src, out, rules, enforce_content_guard=False)
            results = tally(authored, output_texts(Document(str(out))))

            print(f"{name}  (guard changed flag: {result.content_changed})")
            for category, r in results.items():
                category_rows += 1
                string_total += r["total"]
                line = f"  {category}: {r['exact']}/{r['total']} exact"
                if r["guard_equivalent"]:
                    line += f", {r['guard_equivalent']}/{r['total']} guard-equivalent"
                if r["lost"]:
                    line += f", {len(r['lost'])}/{r['total']} LOST"
                print(line)
                for s in r["lost"]:
                    print(f"    lost: {s!r}")
                    lost_total += 1
            print()

    print(f"{len(builders)} documents, {category_rows} category rows, {string_total} authored strings")
    print()
    print("known gaps, not a complete list (not generated, not verified):")
    for item in NOT_COVERED:
        print(f"  {item}")
    print()
    print("headers and footers are written by the formatter on purpose and are")
    print("outside the content guard, so they are not part of this benchmark.")

    if lost_total:
        print()
        print(f"{lost_total} string(s) lost")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
