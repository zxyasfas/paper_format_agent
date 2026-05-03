from __future__ import annotations

import argparse
import json
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .engine import build_simple_report, format_docx, save_report_files
from .llm import LLMConfig, generate_suggestions
from .rules import extract_rules_from_text


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    out.append(c.text)
    return "\n".join(out)


def read_format_text(path: str | Path) -> str:
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    with tempfile.TemporaryDirectory() as td:
        out_dir = Path(td)
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=90,
            )
            converted = out_dir / f"{path.stem}.docx"
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass
    return ""


def main():
    parser = argparse.ArgumentParser(description="Paper Format Agent V2 (safety-first)")
    parser.add_argument("--format-file", required=True, help="格式要求文件（.doc/.docx/.txt）")
    parser.add_argument("--paper-file", required=True, help="论文文件（.docx）")
    parser.add_argument("--out-dir", required=True, help="输出目录")
    parser.add_argument("--use-llm", action="store_true", help="启用 LLM 建议（默认只建议不改文）")
    parser.add_argument("--llm-api-key", default=None, help="LLM API key，默认读取 DEEPSEEK_API_KEY")
    parser.add_argument("--llm-base-url", default="https://api.deepseek.com")
    parser.add_argument("--llm-model", default="deepseek-v4-pro")
    parser.add_argument("--llm-timeout", type=int, default=90)
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    format_text = read_format_text(args.format_file)
    rules = extract_rules_from_text(format_text)
    (out_dir / "format_rules.json").write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

    llm_cfg = LLMConfig(
        enabled=bool(args.use_llm),
        api_key=args.llm_api_key or os.getenv("DEEPSEEK_API_KEY"),
        base_url=args.llm_base_url,
        model=args.llm_model,
        timeout_seconds=args.llm_timeout,
    )
    llm_report = generate_suggestions(args.paper_file, format_text, llm_cfg)
    (out_dir / "llm_suggestions.json").write_text(json.dumps(llm_report, ensure_ascii=False, indent=2), encoding="utf-8")

    output_docx = out_dir / "formatted_paper_v2.docx"
    result = format_docx(args.paper_file, output_docx, rules)
    (out_dir / "modify_log.json").write_text(json.dumps(result.log, ensure_ascii=False, indent=2), encoding="utf-8")

    report = build_simple_report(output_docx, rules, result.removed_numpr_count)
    report["llm_used"] = bool(llm_report.get("used"))
    report["llm_warnings"] = llm_report.get("warnings", [])
    save_report_files(report, out_dir / "format_report.json", out_dir / "format_report.html")

    print(
        json.dumps(
            {
                "output": str(output_docx),
                "score": report["score"],
                "chars_no_space": report["chars_no_space"],
                "removed_numpr_count": report["removed_numpr_count"],
                "llm_used": bool(llm_report.get("used")),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()

