from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from paper_format_agent.ooxml import set_font_east_asia


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def is_abstract_title(text: str) -> bool:
    t = (text or "").strip()
    return bool(re.fullmatch(r"摘\s*要", t))


def is_english_abstract_title(text: str) -> bool:
    return (text or "").strip().upper() == "ABSTRACT"


def is_toc_title(text: str) -> bool:
    return normalize_text(text) in {"目录", "目錄", "目次"}


def is_keyword_zh(text: str) -> bool:
    return bool(re.match(r"^(关键词|关键字)\s*[:：]", (text or "").strip()))


def is_keyword_en(text: str) -> bool:
    return bool(re.match(r"^Keywords?\s*[:：]", (text or "").strip(), flags=re.IGNORECASE))


def is_references(text: str) -> bool:
    return "参考文献" in normalize_text(text)


def is_ack(text: str) -> bool:
    return "致谢" in normalize_text(text) or "致謝" in normalize_text(text)


def is_chapter(text: str) -> bool:
    s = (text or "").strip()
    return bool(re.match(r"^第[一二三四五六七八九十0-9]+章", s)) or bool(re.match(r"^[一二三四五六七八九十]+、", s))


def is_section(text: str) -> bool:
    s = (text or "").strip()
    return bool(re.match(r"^\d+\.\d+\s*", s)) or bool(re.match(r"^第[一二三四五六七八九十]+节", s))


def is_subsection(text: str) -> bool:
    s = (text or "").strip()
    return bool(re.match(r"^\d+\.\d+\.\d+\s*", s)) or bool(re.match(r"^\d+\.\s*", s)) or bool(re.match(r"^（[一二三四五六七八九十]+）", s))


def looks_like_toc_entry(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if re.search(r"[·\.…]{3,}\s*\d+\s*$", t):
        return True
    if is_chapter(t) or is_section(t) or is_subsection(t):
        return True
    return False


def clear_paragraph_numbering(paragraph) -> bool:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)
        return True
    return False


def ensure_style(doc: Document, name: str, font: str, size_pt: float, bold: bool = False, alignment: str | None = None):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    if alignment == "center":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "left":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == "justify":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return style


def set_run_style(paragraph, font: str, size_pt: float, bold: bool | None = None):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = font
        set_font_east_asia(run, font)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_para_format(paragraph, *, align: str, line_spacing: float = 1.25, indent_chars: int = 0, space_before: int = 0, space_after: int = 0):
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = Pt(12 * indent_chars)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def detect_prefix_toc_range(doc: Document) -> tuple[int, int] | None:
    paras = list(doc.paragraphs)
    idx_abs = None
    for i, p in enumerate(paras):
        if is_abstract_title((p.text or "").strip()):
            idx_abs = i
            break
    if idx_abs is None or idx_abs < 10:
        return None

    start = 0
    while start < idx_abs and not (paras[start].text or "").strip():
        start += 1
    if start >= idx_abs:
        return None

    toc_like = 0
    total = 0
    for p in paras[start:idx_abs]:
        t = (p.text or "").strip()
        if not t:
            continue
        total += 1
        if looks_like_toc_entry(t) and len(normalize_text(t)) <= 45:
            toc_like += 1

    if total >= 8 and toc_like / total >= 0.7:
        return (start, idx_abs)
    return None


def set_page_setup(doc: Document, rules: dict):
    m = rules["margins_cm"]
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(m["top"])
        section.bottom_margin = Cm(m["bottom"])
        section.left_margin = Cm(m["left"])
        section.right_margin = Cm(m["right"])


def set_defaults(doc: Document, rules: dict):
    normal = doc.styles["Normal"]
    normal.font.name = rules["body"]["font"]
    normal.font.size = Pt(rules["body"]["size_pt"])
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), rules["body"]["font"])
    rfonts.set(qn("w:ascii"), rules["body"]["font"])
    rfonts.set(qn("w:hAnsi"), rules["body"]["font"])

    ensure_style(doc, "Title", rules["toc"]["font"], rules["toc"]["title_size_pt"], True, "center")
    ensure_style(doc, "Heading 1", rules["heading_1"]["font"], rules["heading_1"]["size_pt"], True, "center")
    ensure_style(doc, "Heading 2", rules["heading_2"]["font"], rules["heading_2"]["size_pt"], True, "center")
    ensure_style(doc, "Heading 3", rules["heading_3"]["font"], rules["heading_3"]["size_pt"], True, "left")


@dataclass
class FormatResult:
    output: str
    chars_no_space: int
    removed_numpr_count: int
    log: list[dict[str, Any]] = field(default_factory=list)


def count_chars(doc: Document) -> int:
    joined = []
    for p in doc.paragraphs:
        if p.text:
            joined.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    joined.append(c.text)
    return len(re.sub(r"\s+", "", "\n".join(joined)))


def format_docx(input_docx: str | Path, output_docx: str | Path, rules: dict) -> FormatResult:
    doc = Document(str(input_docx))
    set_defaults(doc, rules)
    set_page_setup(doc, rules)

    prefix_toc = detect_prefix_toc_range(doc)
    removed_numpr = 0
    log: list[dict[str, Any]] = []

    if prefix_toc is not None:
        log.append({"action": "detect_prefix_toc", "range": list(prefix_toc), "note": "检测到摘要前手工目录块，按目录正文样式处理，不改文本内容。"})

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if clear_paragraph_numbering(p):
            removed_numpr += 1
        p.paragraph_format.page_break_before = False
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        in_prefix_toc = prefix_toc is not None and prefix_toc[0] <= i < prefix_toc[1]
        if in_prefix_toc:
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["toc"]["font"], rules["toc"]["body_size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=0, space_after=0)
            continue

        if is_abstract_title(text):
            p.style = doc.styles["Title"]
            set_run_style(p, rules["abstract_title"]["font"], rules["abstract_title"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_english_abstract_title(text):
            p.style = doc.styles["Title"]
            set_run_style(p, rules["english"]["font"], 18, True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_toc_title(text):
            p.style = doc.styles["Title"]
            set_run_style(p, rules["toc"]["font"], rules["toc"]["title_size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_references(text) or is_ack(text):
            p.style = doc.styles["Title"]
            set_run_style(p, rules["heading_1"]["font"], rules["heading_1"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_chapter(text):
            p.style = doc.styles["Heading 1"]
            set_run_style(p, rules["heading_1"]["font"], rules["heading_1"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=12, space_after=12)
        elif is_section(text):
            p.style = doc.styles["Heading 2"]
            set_run_style(p, rules["heading_2"]["font"], rules["heading_2"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=6, space_after=6)
        elif is_subsection(text):
            p.style = doc.styles["Heading 3"]
            set_run_style(p, rules["heading_3"]["font"], rules["heading_3"]["size_pt"], True)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=6, space_after=6)
        elif is_keyword_zh(text):
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["body"]["font"], rules["body"]["size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=6, space_after=12)
        elif is_keyword_en(text):
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["english"]["font"], rules["english"]["size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=6, space_after=12)
        else:
            p.style = doc.styles["Normal"]
            font = rules["english"]["font"] if re.search(r"[A-Za-z]{6,}", text) and not re.search(r"[\u4e00-\u9fff]", text) else rules["body"]["font"]
            set_run_style(p, font, rules["body"]["size_pt"], False)
            set_para_format(
                p,
                align="justify",
                line_spacing=rules["body"]["line_spacing"],
                indent_chars=rules["body"]["first_line_indent_chars"],
                space_before=0,
                space_after=0,
            )

    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return FormatResult(
        output=str(output_docx),
        chars_no_space=count_chars(doc),
        removed_numpr_count=removed_numpr,
        log=log,
    )


def build_simple_report(docx_path: str | Path, rules: dict, removed_numpr_count: int) -> dict[str, Any]:
    doc = Document(str(docx_path))
    chars = count_chars(doc)
    has_abs = any(is_abstract_title((p.text or "").strip()) for p in doc.paragraphs)
    has_en_abs = any(is_english_abstract_title((p.text or "").strip()) for p in doc.paragraphs)
    has_kw_zh = any(is_keyword_zh((p.text or "").strip()) for p in doc.paragraphs)
    has_kw_en = any(is_keyword_en((p.text or "").strip()) for p in doc.paragraphs)
    numpr_left = sum(1 for p in doc.paragraphs if "w:numPr" in p._p.xml)
    score = 60
    score += 10 if has_abs else 0
    score += 5 if has_kw_zh else 0
    score += 5 if has_en_abs else 0
    score += 5 if has_kw_en else 0
    score += 5 if numpr_left == 0 else 0
    score += 10 if chars >= int(rules.get("min_total_chars_no_space", 10000)) else 0
    score = min(score, 100)
    return {
        "score": float(score),
        "chars_no_space": chars,
        "has_abstract": has_abs,
        "has_english_abstract": has_en_abs,
        "has_keywords_zh": has_kw_zh,
        "has_keywords_en": has_kw_en,
        "removed_numpr_count": removed_numpr_count,
        "remaining_numpr_count": numpr_left,
    }


def save_report_files(report: dict[str, Any], out_json: Path, out_html: Path):
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    html = (
        "<!doctype html><html><head><meta charset='utf-8'><title>V2 格式报告</title>"
        "<style>body{font-family:Arial,'Microsoft YaHei';max-width:960px;margin:24px auto;line-height:1.6}"
        "table{border-collapse:collapse;width:100%}td,th{border:1px solid #ddd;padding:8px}"
        "th{background:#f5f5f5}</style></head><body>"
        f"<h1>V2 格式报告</h1><p><b>评分：</b>{report['score']}</p>"
        "<table><tr><th>项</th><th>值</th></tr>"
        f"<tr><td>不含空白字符数</td><td>{report['chars_no_space']}</td></tr>"
        f"<tr><td>中文摘要</td><td>{report['has_abstract']}</td></tr>"
        f"<tr><td>英文摘要</td><td>{report['has_english_abstract']}</td></tr>"
        f"<tr><td>中文关键词</td><td>{report['has_keywords_zh']}</td></tr>"
        f"<tr><td>英文关键词</td><td>{report['has_keywords_en']}</td></tr>"
        f"<tr><td>清除编号段落数</td><td>{report['removed_numpr_count']}</td></tr>"
        f"<tr><td>剩余编号段落数</td><td>{report['remaining_numpr_count']}</td></tr>"
        "</table></body></html>"
    )
    out_html.write_text(html, encoding="utf-8")

