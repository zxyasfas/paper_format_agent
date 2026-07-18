from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import paper_format_agent.pipeline as pipeline_module
from paper_format_agent.pipeline import run_pipeline
from paper_format_agent.rules import extract_rules_from_text


class PipelineContentGuardTests(unittest.TestCase):
    def test_pipeline_keeps_content_fingerprint(self):
        with TemporaryDirectory() as td:
            td_path = Path(td)
            src = td_path / "input.docx"
            out = td_path / "out.docx"

            doc = Document()
            doc.add_paragraph("摘要")
            doc.add_paragraph("这是摘要内容。")
            doc.add_paragraph("关键词：测试；格式化")
            doc.add_paragraph("一、绪论")
            doc.add_paragraph("这是正文第一段。")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "列1"
            table.cell(0, 1).text = "列2"
            doc.save(src)

            rules = extract_rules_from_text("正文宋体小四，1.25倍行距。")
            result = run_pipeline(src, out, rules, enforce_content_guard=True)
            self.assertFalse(result.content_changed)
            self.assertTrue(out.exists())

    def test_pipeline_writes_header_and_page_number_footer(self):
        with TemporaryDirectory() as td:
            td_path = Path(td)
            src = td_path / "input.docx"
            out = td_path / "out.docx"

            doc = Document()
            doc.add_paragraph("Abstract")
            doc.add_paragraph("Body paragraph.")
            doc.save(src)

            rules = extract_rules_from_text("")
            rules["header"]["text"] = "Journal Draft"
            result = run_pipeline(src, out, rules, enforce_content_guard=True)

            formatted = Document(str(out))
            section = formatted.sections[0]
            self.assertEqual(section.header.paragraphs[0].text, "Journal Draft")
            self.assertIn("PAGE", section.footer._element.xml)
            self.assertFalse(result.content_changed)
            self.assertEqual(result.logs[0]["action"], "setup_document_base")
            self.assertEqual(result.logs[0]["headers_written"], 1)
            self.assertEqual(result.logs[0]["footers_written"], 1)

    def test_pipeline_keeps_reference_spacing_and_indent_stable(self):
        with TemporaryDirectory() as td:
            td_path = Path(td)
            src = td_path / "input.docx"
            out = td_path / "out.docx"

            doc = Document()
            doc.add_paragraph("References")
            doc.add_paragraph(
                "Smith, J. Q. (2024). Synthetic testing for reference layouts. "
                "Journal of Fake Results, 12(3), 45-67."
            )
            doc.add_paragraph(
                "Lopez, M., & Chen, A. (2023). Paragraph spacing in mock APA samples. "
                "Example Review, 8(1), 10-18."
            )
            doc.save(src)

            rules = extract_rules_from_text("正文宋体小四，1.25倍行距。")
            result = run_pipeline(src, out, rules, enforce_content_guard=True)

            formatted = Document(str(out))
            self.assertEqual(formatted.paragraphs[0].text, "References")

            for reference_paragraph in formatted.paragraphs[1:]:
                self.assertEqual(reference_paragraph.style.name, "Normal")
                self.assertEqual(reference_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
                self.assertEqual(reference_paragraph.paragraph_format.line_spacing, 1.25)
                self.assertEqual(reference_paragraph.paragraph_format.first_line_indent.pt, 0.0)
                self.assertEqual(reference_paragraph.paragraph_format.space_before.pt, 0.0)
                self.assertEqual(reference_paragraph.paragraph_format.space_after.pt, 0.0)

            self.assertFalse(result.content_changed)

    def test_pipeline_aborts_when_styling_step_changes_text(self):
        with TemporaryDirectory() as td:
            td_path = Path(td)
            src = td_path / "input.docx"
            out = td_path / "out.docx"

            doc = Document()
            doc.add_paragraph("一、绪论")
            doc.add_paragraph("这是正文第一段。")
            doc.save(src)

            real_styles = pipeline_module.apply_final_styles_from_markers

            def styles_that_also_edit_text(document, rules):
                count = real_styles(document, rules)
                for p in document.paragraphs:
                    if "正文第一段" in p.text:
                        p.text = p.text.replace("第一段", "第1段")
                        break
                return count

            rules = extract_rules_from_text("正文宋体小四，1.25倍行距。")
            with patch.object(
                pipeline_module,
                "apply_final_styles_from_markers",
                styles_that_also_edit_text,
            ):
                with self.assertRaisesRegex(
                    ValueError,
                    r"^content guard failed: non-whitespace content changed$",
                ):
                    run_pipeline(src, out, rules, enforce_content_guard=True)

            self.assertFalse(out.exists())


if __name__ == "__main__":
    unittest.main()
