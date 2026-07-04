from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from paper_format_agent.service import format_paper, read_format_text


SYNTHETIC_GUIDE = """# Synthetic Formatting Guide

## Page Setup
- Paper size: A4.

## Body Text
- Chinese body text: Songti, 12 pt, 1.25 line spacing.
- First-line indent: 2 Chinese characters.

## Sections
- Chinese abstract is required.
- Chinese keywords are required.
"""


def _make_paper(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("合成论文标题")
    doc.add_paragraph("摘要")
    doc.add_paragraph(
        "这是一段用于测试的合成正文文本，长度足够被判定为正文段落，"
        "不涉及任何真实数据，仅用来验证内容保护与格式化流程。"
    )
    doc.add_paragraph("关键词：测试；合成；格式")
    doc.add_paragraph("这是正文第二段，同样是合成文本，用于验证排版前后内容一致。")
    doc.save(path)


class FormatPaperServiceTests(unittest.TestCase):
    def test_format_paper_preserves_content_and_reports_fingerprints(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            guide = td / "guide.txt"
            guide.write_text(SYNTHETIC_GUIDE, encoding="utf-8")
            paper = td / "paper.docx"
            _make_paper(paper)
            out_dir = td / "out"

            result = format_paper(
                paper,
                out_dir,
                format_file=guide,
                engine="python",
                strict_required_sections=False,
            )

            # Content guard: body text unchanged, fingerprints match, guard enforced.
            self.assertFalse(result["content_changed"])
            self.assertTrue(result["content_guard_enforced"])
            self.assertEqual(
                result["content_fingerprint_before"],
                result["content_fingerprint_after"],
            )
            # The formatted document and its report are actually written.
            self.assertTrue(Path(result["output"]).exists())
            self.assertTrue((out_dir / "format_report.json").exists())

    def test_format_paper_requires_a_rule_source(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            paper = td / "paper.docx"
            _make_paper(paper)
            with self.assertRaises(ValueError):
                format_paper(paper, td / "out")  # no rules / format_text / format_file

    def test_read_format_text_reads_txt(self):
        with tempfile.TemporaryDirectory() as td:
            guide = Path(td) / "guide.txt"
            guide.write_text(SYNTHETIC_GUIDE, encoding="utf-8")
            text = read_format_text(guide)
            self.assertIn("Body Text", text)


if __name__ == "__main__":
    unittest.main()
