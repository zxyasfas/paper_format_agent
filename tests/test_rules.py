from __future__ import annotations

import unittest

from docx import Document

from paper_format_agent.pipeline import T_FIG_CAPTION, T_H1, T_H2, T_TABLE_CAPTION, classify_document
from paper_format_agent.rules import extract_rules_from_text


class RulesExtractionTests(unittest.TestCase):
    def test_extract_core_rules(self):
        text = (
            "页边距：上2.54cm，下2.54cm，左3.17cm，右2.54cm。"
            "正文宋体小四，1.25倍行距。"
            "包含目录、英文摘要、英文关键词。"
            "字数不少于1万字。"
            "一级标题居中。"
        )
        rules = extract_rules_from_text(text)
        self.assertEqual(rules["margins_cm"]["left"], 3.17)
        self.assertEqual(rules["margins_cm"]["right"], 2.54)
        self.assertEqual(rules["body"]["size_pt"], 12.0)
        self.assertEqual(rules["body"]["line_spacing"], 1.25)
        self.assertTrue(rules["required_sections"]["toc"])
        self.assertTrue(rules["required_sections"]["en_abstract"])
        self.assertTrue(rules["required_sections"]["en_keywords"])
        self.assertEqual(rules["min_total_chars_no_space"], 10000)
        self.assertEqual(rules["heading_1"]["align"], "center")

    def test_ieee_style_headings_and_captions_are_classified(self):
        doc = Document()
        doc.add_paragraph("I. INTRODUCTION")
        doc.add_paragraph("A. Preparation of Papers")
        doc.add_paragraph("Fig. 1. System overview")
        doc.add_paragraph("Table I. Experimental settings")

        cls = classify_document(doc)
        self.assertEqual(cls.types[0], T_H1)
        self.assertEqual(cls.types[1], T_H2)
        self.assertEqual(cls.types[2], T_FIG_CAPTION)
        self.assertEqual(cls.types[3], T_TABLE_CAPTION)


if __name__ == "__main__":
    unittest.main()
