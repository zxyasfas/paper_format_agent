from __future__ import annotations

import unittest
from tempfile import TemporaryDirectory
from pathlib import Path

from docx import Document

from paper_format_agent.quality import build_content_fingerprint, normalize_for_content_guard


class QualityGuardTests(unittest.TestCase):
    def test_normalize_ignores_space_and_bullets(self):
        a = "  关键词：A B C  "
        b = "▪关键词：ABC"
        self.assertEqual(normalize_for_content_guard(a), normalize_for_content_guard(b))

    def test_fingerprint_stable_for_equivalent_content(self):
        with TemporaryDirectory() as td:
            p1 = Path(td) / "a.docx"
            p2 = Path(td) / "b.docx"
            d1 = Document()
            d1.add_paragraph("摘要")
            d1.add_paragraph("  这是正文。")
            d1.save(p1)

            d2 = Document()
            d2.add_paragraph("摘要")
            d2.add_paragraph("▪这是正文。")
            d2.save(p2)

            f1 = build_content_fingerprint(Document(str(p1)))
            f2 = build_content_fingerprint(Document(str(p2)))
            self.assertEqual(f1, f2)


if __name__ == "__main__":
    unittest.main()

