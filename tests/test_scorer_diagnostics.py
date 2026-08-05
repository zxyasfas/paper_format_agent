from __future__ import annotations

import json
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from paper_format_agent.scorer import save_reports, score_document


class ScorerDiagnosticsTests(unittest.TestCase):
    def test_score_document_returns_actionable_diagnostics(self):
        with TemporaryDirectory() as td:
            docx_path = Path(td) / "paper.docx"
            doc = Document()
            doc.add_paragraph("Body paragraph without required front matter.")
            doc.save(docx_path)

            report = score_document(
                docx_path,
                {
                    "required_sections": {
                        "zh_abstract": True,
                        "zh_keywords": True,
                        "toc": True,
                    },
                    "min_total_chars_no_space": 0,
                },
                enforce_required_sections=True,
            )

            diagnostics = {item["name"]: item for item in report["diagnostics"]}
            self.assertIn("missing_zh_abs", diagnostics)
            self.assertIn("missing_zh_keywords", diagnostics)
            self.assertIn("missing_toc_title", diagnostics)
            self.assertEqual(diagnostics["missing_zh_abs"]["severity"], "high")
            self.assertTrue(diagnostics["missing_zh_abs"]["suggested_fix"])
            self.assertEqual(
                diagnostics["missing_zh_abs"]["penalty"],
                25,
            )

    def test_save_reports_writes_diagnostics_to_json_and_html(self):
        with TemporaryDirectory() as td:
            out_json = Path(td) / "format_report.json"
            out_html = Path(td) / "format_report.html"
            report = {
                "score": 75,
                "raw_quality_score": 75,
                "chars_no_space": 100,
                "features": {},
                "penalties": [{"name": "missing_zh_abs", "value": 25}],
                "diagnostics": [
                    {
                        "name": "missing_zh_abs",
                        "category": "required_sections",
                        "severity": "high",
                        "penalty": 25,
                        "summary": "Chinese abstract section is missing.",
                        "suggested_fix": "Add a standalone Chinese abstract title.",
                        "evidence": {},
                    }
                ],
            }

            save_reports(report, out_json, out_html)

            saved = json.loads(out_json.read_text(encoding="utf-8"))
            html = out_html.read_text(encoding="utf-8")
            self.assertEqual(saved["diagnostics"][0]["name"], "missing_zh_abs")
            self.assertIn("Actionable diagnostics", html)
            self.assertIn("Suggested fix", html)
            self.assertIn("missing_zh_abs", html)

    def test_table_without_caption_triggers_missing_table_caption(self):
        with TemporaryDirectory() as td:
            docx_path = Path(td) / "table_no_caption.docx"
            doc = Document()
            doc.add_paragraph("The experimental results are shown below.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "method"
            table.cell(0, 1).text = "accuracy"
            doc.save(docx_path)

            report = score_document(docx_path, {"min_total_chars_no_space": 0})
            diagnostics = {item["name"]: item for item in report["diagnostics"]}

            self.assertIn("missing_table_caption", diagnostics)
            diag = diagnostics["missing_table_caption"]
            self.assertEqual(diag["severity"], "medium")
            self.assertTrue(diag["suggested_fix"])
            self.assertGreaterEqual(diag["penalty"], 6)
            self.assertGreaterEqual(report["penalties"][0]["value"], 6)
            evidence = diag["evidence"]["table_caption_evidence"]
            self.assertEqual(len(evidence), 1)
            self.assertEqual(evidence[0]["table_index"], 0)
            self.assertIn("idx", evidence[0])
            self.assertIn("text", evidence[0])

    def test_table_with_caption_does_not_trigger_diagnostic(self):
        with TemporaryDirectory() as td:
            base_rules = {"min_total_chars_no_space": 0}
            captioned_path = Path(td) / "captioned.docx"
            bare_path = Path(td) / "bare.docx"

            captioned = Document()
            captioned.add_paragraph("Table 1: Experimental results")
            table = captioned.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "method"
            table.cell(0, 1).text = "accuracy"
            captioned.save(captioned_path)

            bare = Document()
            bare.add_paragraph("The experimental results are shown below.")
            table = bare.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "method"
            table.cell(0, 1).text = "accuracy"
            bare.save(bare_path)

            captioned_report = score_document(captioned_path, base_rules)
            bare_report = score_document(bare_path, base_rules)

            names = {item["name"] for item in captioned_report["diagnostics"]}
            self.assertNotIn("missing_table_caption", names)
            self.assertNotIn("missing_table_caption", captioned_report["penalties"])
            self.assertGreater(captioned_report["score"], bare_report["score"])


if __name__ == "__main__":
    unittest.main()
