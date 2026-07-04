"""Reusable formatting service.

This module holds the format-a-single-paper pipeline as a plain function so it can
be called the same way from the CLI, the GUI, and the MCP server, without any of
them re-implementing the orchestration. Keeping one code path here is what lets the
content-guard behaviour stay identical across every entry point.
"""

from __future__ import annotations

import json
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .engines import run_postprocess_engine
from .llm import LLMConfig, generate_suggestions
from .pipeline import run_pipeline
from .rules import extract_rules_from_text
from .scorer import save_reports, score_document


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    out.append(c.text)
    return "\n".join(out)


def read_format_text(path: str | Path) -> str:
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    converted_adjacent = path.with_name(path.stem + "_converted.docx")
    if converted_adjacent.exists():
        return read_docx_text(converted_adjacent)

    with tempfile.TemporaryDirectory() as td:
        out_dir = Path(td)
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=90,
            )
            converted = out_dir / f"{path.stem}.docx"
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass

        # Fallback: Word COM conversion on Windows.
        try:
            converted = out_dir / f"{path.stem}.docx"
            ps = rf"""
$ErrorActionPreference = "Stop"
$src = "{str(path.resolve())}"
$dst = "{str(converted.resolve())}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($src, $false, $true)
$doc.SaveAs([ref]$dst, [ref]16)
$doc.Close()
$word.Quit()
"""
            subprocess.run(
                ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=120,
            )
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass
    return ""


def format_paper(
    paper_file: str | Path,
    out_dir: str | Path,
    *,
    format_file: str | Path | None = None,
    format_text: str | None = None,
    rules: dict | None = None,
    engine: str = "auto",
    strict_required_sections: bool = False,
    allow_content_change: bool = False,
    marker_dump: bool = False,
    calibration_file: str | Path | None = None,
    use_llm: bool = False,
    llm_api_key: str | None = None,
    llm_base_url: str = "https://api.deepseek.com",
    llm_model: str = "deepseek-v4-pro",
    llm_timeout: int = 90,
) -> dict:
    """Format one paper and write the full artifact set into ``out_dir``.

    Provide the formatting rules one of three ways: ``rules`` (already extracted),
    ``format_text`` (raw guide text), or ``format_file`` (a .docx/.doc/.txt guide
    that will be read and parsed). The content guard is enforced unless
    ``allow_content_change`` is set, in which case a changed body fingerprint is
    permitted instead of raising.

    Returns a summary dict; the machine-readable ``format_report.json`` written to
    ``out_dir`` carries the full report including both content fingerprints.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    paper_file = Path(paper_file)

    if rules is None:
        if format_text is None:
            if format_file is None:
                raise ValueError("format_paper requires one of: rules, format_text, or format_file")
            format_text = read_format_text(format_file)
        rules = extract_rules_from_text(format_text)
    if format_text is None:
        format_text = ""

    (out_dir / "format_rules.json").write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

    llm_cfg = LLMConfig(
        enabled=bool(use_llm),
        api_key=llm_api_key or os.getenv("DEEPSEEK_API_KEY"),
        base_url=llm_base_url,
        model=llm_model,
        timeout_seconds=llm_timeout,
    )
    llm_report = generate_suggestions(paper_file, format_text, llm_cfg)
    (out_dir / "llm_suggestions.json").write_text(json.dumps(llm_report, ensure_ascii=False, indent=2), encoding="utf-8")

    marker_dump_path = out_dir / "marker_dump.json" if marker_dump else None
    output_docx = out_dir / "formatted_paper_v3.docx"
    run_result = run_pipeline(
        paper_file,
        output_docx,
        rules,
        write_marker_dump=marker_dump_path,
        enforce_content_guard=not bool(allow_content_change),
    )
    (out_dir / "modify_log.json").write_text(json.dumps(run_result.logs, ensure_ascii=False, indent=2), encoding="utf-8")

    engine_report = run_postprocess_engine(engine, output_docx)
    (out_dir / "engine_report.json").write_text(json.dumps(engine_report, ensure_ascii=False, indent=2), encoding="utf-8")

    report_before = score_document(
        paper_file,
        rules,
        calibration_file=calibration_file,
        baseline_docx=paper_file,
        enforce_required_sections=bool(strict_required_sections),
    )
    report_after = score_document(
        output_docx,
        rules,
        calibration_file=calibration_file,
        baseline_docx=paper_file,
        enforce_required_sections=bool(strict_required_sections),
    )

    report = report_after.copy()
    report["score_before"] = round(report_before["score"], 1)
    report["score_after"] = round(report_after["score"], 1)
    report["score_improvement"] = round(report_after["score"] - report_before["score"], 1)
    report["chars_no_space_before"] = report_before["chars_no_space"]
    report["chars_no_space_after"] = report_after["chars_no_space"]
    report["llm_used"] = bool(llm_report.get("used"))
    report["llm_warnings"] = llm_report.get("warnings", [])
    report["engine_report"] = engine_report
    report["removed_numpr_count"] = run_result.removed_numpr_count
    report["classification_confidence"] = run_result.classification_confidence
    report["content_fingerprint_before"] = run_result.content_fingerprint_before
    report["content_fingerprint_after"] = run_result.content_fingerprint_after
    report["content_changed"] = bool(run_result.content_changed)
    report["content_guard_enforced"] = not bool(allow_content_change)
    save_reports(report, out_dir / "format_report.json", out_dir / "format_report.html")

    return {
        "paper_file": str(paper_file),
        "out_dir": str(out_dir),
        "output": str(output_docx),
        "score_before": report["score_before"],
        "score_after": report["score_after"],
        "score_improvement": report["score_improvement"],
        "raw_quality_score": report.get("raw_quality_score"),
        "chars_no_space_before": report["chars_no_space_before"],
        "chars_no_space_after": report["chars_no_space_after"],
        "removed_numpr_count": run_result.removed_numpr_count,
        "content_changed": bool(run_result.content_changed),
        "content_guard_enforced": not bool(allow_content_change),
        "content_fingerprint_before": run_result.content_fingerprint_before,
        "content_fingerprint_after": run_result.content_fingerprint_after,
        "engine": engine_report.get("engine", engine),
        "engine_success": bool(engine_report.get("success")),
        "llm_used": bool(llm_report.get("used")),
    }
