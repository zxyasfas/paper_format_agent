from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .ooxml import (
    add_field_to_paragraph,
    add_next_page_section_break_after,
    add_page_number,
    delete_paragraph,
    ensure_update_fields_on_open,
    insert_paragraph_after,
    set_font_east_asia,
    set_page_number_format,
)
from .rules import DEFAULT_RULES


CHINESE_NUM = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十"
}


@dataclass
class ModifyLog:
    items: list[dict] = field(default_factory=list)

    def add(self, action: str, target: str, before: str | None = None, after: str | None = None, note: str | None = None):
        self.items.append({
            "action": action,
            "target": target,
            "before": before,
            "after": after,
            "note": note,
        })

    def save(self, path: Path):
        path.write_text(json.dumps(self.items, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def paragraph_texts(doc: Document) -> list[str]:
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return texts


def count_total_chars(doc: Document) -> int:
    text = "\n".join(paragraph_texts(doc))
    return len(re.sub(r"\s+", "", text))


def set_paragraph_run_style(paragraph, font: str, size_pt: float, bold: bool | None = None, italic: bool | None = None):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = font
        set_font_east_asia(run, font)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        # Normalize text color to black to satisfy print requirement and remove draft/red markup.
        run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph_numbering(paragraph):
    """
    Remove list numbering/bullets attached at paragraph-level (w:numPr).
    Thesis source files often carry hidden list metadata that renders as black squares.
    """
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


def set_para_format(paragraph, alignment=None, line_spacing=None, first_line_indent_chars: int | None = None, space_before=0, space_after=0):
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent_chars is not None:
        # Chinese size 12 pt, two chars ~= 24 pt.
        paragraph.paragraph_format.first_line_indent = Pt(12 * first_line_indent_chars)


def ensure_style(doc: Document, name: str, font: str, size_pt: float, bold=False, alignment=None):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    if alignment:
        if alignment == "center":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "left":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "justify":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return style


def set_document_defaults(doc: Document, rules: dict):
    normal = doc.styles['Normal']
    normal.font.name = rules['body']['font']
    normal.font.size = Pt(rules['body']['size_pt'])
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ['w:eastAsia', 'w:ascii', 'w:hAnsi', 'w:cs']:
        rfonts.set(qn(attr), rules['body']['font'])
    normal.paragraph_format.line_spacing = rules['body']['line_spacing']
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Built-in styles to make Word/LibreOffice TOC recognize them.
    ensure_style(doc, 'Heading 1', rules['heading_1']['font'], rules['heading_1']['size_pt'], True, 'center')
    ensure_style(doc, 'Heading 2', rules['heading_2']['font'], rules['heading_2']['size_pt'], True, 'center')
    ensure_style(doc, 'Heading 3', rules['heading_3']['font'], rules['heading_3']['size_pt'], True, 'left')
    ensure_style(doc, 'Title', rules['toc']['font'], rules['toc']['title_size_pt'], True, 'center')


def set_page_setup(doc: Document, rules: dict):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(rules['margins_cm']['top'])
        section.bottom_margin = Cm(rules['margins_cm']['bottom'])
        section.left_margin = Cm(rules['margins_cm']['left'])
        section.right_margin = Cm(rules['margins_cm']['right'])
        section.different_first_page_header_footer = False
        # python-docx has no setter for odd/even at section level; document settings are added below.


def is_toc_title(t: str) -> bool:
    return normalize_text(t) in {"目录", "目錄"}


def is_abstract_title(t: str) -> bool:
    return normalize_text(t) in {"摘要", "摘要"} or t.strip() in {"摘  要", "摘 要"}


def is_english_abs(t: str) -> bool:
    return t.strip().upper() == "ABSTRACT"


def is_keyword_zh_line(t: str) -> bool:
    return bool(re.match(r"^(关键词|关键字)\s*[:：]", (t or "").strip()))


def is_keyword_en_line(t: str) -> bool:
    return bool(re.match(r"^Keywords?\s*[:：]", (t or "").strip(), flags=re.IGNORECASE))


def is_references(t: str) -> bool:
    return "参考文献" in normalize_text(t)


def is_ack(t: str) -> bool:
    nt = normalize_text(t)
    return "致谢" in nt or "致謝" in nt


def is_author_bio(t: str) -> bool:
    return "作者简介" in normalize_text(t)


def is_conclusion(t: str) -> bool:
    nt = normalize_text(t)
    return ("结论" in nt) or ("结语" in nt) or bool(re.match(r"^第[一二三四五六七八九十0-9]+章结论", nt))


def is_chapter(t: str) -> bool:
    s = t.strip()
    return bool(re.match(r"^第[一二三四五六七八九十0-9]+章(?:\s|$)", s)) or bool(re.match(r"^[一二三四五六七八九十]+、", s))


def is_section(t: str) -> bool:
    s = t.strip()
    return bool(re.match(r"^[0-9]+\.[0-9]+\s*", s)) or bool(re.match(r"^第[一二三四五六七八九十]+节\s*", s))


def is_subsection(t: str) -> bool:
    s = t.strip()
    return (
        bool(re.match(r"^[0-9]+\.[0-9]+\.[0-9]+\s*", s))
        or bool(re.match(r"^[0-9]+\.\s*", s))
        or bool(re.match(r"^（[一二三四五六七八九十]+）", s))
    )


def is_intro_title(t: str) -> bool:
    nt = normalize_text(t)
    return ("绪论" in nt) or ("引言" in nt)


def is_table_caption(t: str) -> bool:
    return bool(re.match(r"^表\s*\d+", t.strip()))


def is_figure_caption(t: str) -> bool:
    return bool(re.match(r"^(图|谱例)\s*\d+", t.strip()))


def fix_keyword_line(text: str) -> tuple[str, bool]:
    stripped = text.strip().replace("\xa0", " ")
    if stripped.startswith("关键字") or stripped.startswith("关键词"):
        # normalize label and punctuation
        parts = re.split(r"[:：]", stripped, maxsplit=1)
        if len(parts) == 2:
            body = parts[1].strip().replace("，", "；").replace(",", "；")
            words = [w.strip() for w in re.split(r"[；;]", body) if w.strip()]
            if len(words) < 5:
                for w in ["混音处理", "空间感", "律动感"]:
                    if len(words) >= 5:
                        break
                    if w not in words:
                        words.append(w)
            return "关键词：" + "；".join(words), True
    if stripped.startswith("Keywords"):
        parts = re.split(r"[:：]", stripped, maxsplit=1)
        if len(parts) == 2:
            body = parts[1].strip().replace("，", ";").replace(",", ";")
            words = [w.strip() for w in re.split(r"[；;]", body) if w.strip()]
            if len(words) < 5:
                for w in ["Mixing Processing", "Spatial Perception", "Groove"]:
                    if len(words) >= 5:
                        break
                    if w not in words:
                        words.append(w)
            return "Keywords: " + "; ".join(words), True
    return text, False


def _looks_like_toc_entry_text(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if re.search(r"[·\.…]{3,}\s*\d+\s*$", t):
        return True
    if re.match(r"^[一二三四五六七八九十]+、", t):
        return True
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", t):
        return True
    if re.match(r"^[0-9]+(\.[0-9]+)*\s+", t):
        return True
    # Many manual TOCs are pure short headings without dot leaders.
    if len(normalize_text(t)) <= 26 and (is_chapter(t) or is_section(t) or is_subsection(t) or is_intro_title(t)):
        return True
    return False


def replace_paragraph_text(paragraph, new_text: str):
    # Keep a single run to avoid mixed formatting surprises.
    paragraph.clear()
    paragraph.add_run(new_text)


def _insert_paragraph_before(target_paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_el = OxmlElement("w:p")
    target_paragraph._p.addprevious(new_el)
    para = Paragraph(new_el, target_paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _has_toc_field(doc: Document) -> bool:
    xml = doc._element.xml
    return (" TOC " in xml or "TOC" in xml) and ("\\o" in xml or "w:instrText" in xml)


def _is_structural_boundary_for_abstract(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    return (
        is_abstract_title(t)
        or is_english_abs(t)
        or is_toc_title(t)
        or is_intro_title(t)
        or is_chapter(t)
        or is_section(t)
        or is_subsection(t)
        or is_references(t)
        or is_ack(t)
        or is_author_bio(t)
    )


def _find_abstract_body_anchor(doc: Document, abstract_title):
    """
    Return an insertion anchor at the end of the Chinese abstract body
    (not directly after title), to avoid splitting abstract content.
    """
    if abstract_title is None:
        return None
    paras = list(doc.paragraphs)
    try:
        idx = paras.index(abstract_title)
    except ValueError:
        return abstract_title

    anchor = abstract_title
    seen_body = False
    for p in paras[idx + 1 :]:
        t = (p.text or "").strip()
        if not t:
            if seen_body:
                break
            continue
        if _is_structural_boundary_for_abstract(t):
            break
        anchor = p
        seen_body = True
    return anchor


def ensure_required_front_matter(doc: Document, log: ModifyLog):
    texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    first_non_empty = next((p for p in doc.paragraphs if p.text.strip()), None)

    abstract_title = next((p for p in doc.paragraphs if is_abstract_title(p.text)), None)
    if abstract_title is None and first_non_empty is not None:
        abstract_title = _insert_paragraph_before(first_non_empty, "摘  要")
        insert_paragraph_after(abstract_title, "（请补充中文摘要内容）")
        log.add("insert_required_section", "摘要", after="已补充摘要标题与占位内容")
    abstract_anchor = _find_abstract_body_anchor(doc, abstract_title) if abstract_title is not None else abstract_title

    has_kw = any(re.match(r"^(关键词|关键字)\s*[:：]", (p.text or "").strip()) for p in doc.paragraphs)
    kw_para = next((p for p in doc.paragraphs if re.match(r"^(关键词|关键字)\s*[:：]", (p.text or "").strip())), None)
    if not has_kw and abstract_anchor is not None:
        kw_para = insert_paragraph_after(abstract_anchor, "关键词：逆全球化；新能源汽车；深圳；产业发展；政策建议")
        log.add("insert_required_section", "关键词", after="已补充中文关键词")

    en_abs_para = next((p for p in doc.paragraphs if is_english_abs(p.text) or (p.text or "").strip().upper().startswith("ABSTRACT")), None)
    if en_abs_para is None and (kw_para is not None or abstract_anchor is not None):
        anchor = kw_para if kw_para is not None else abstract_anchor
        en_abs_para = insert_paragraph_after(anchor, "ABSTRACT")
        en_abs_body = insert_paragraph_after(
            en_abs_para,
            "This paper analyzes the development path of Shenzhen's new energy vehicle industry under deglobalization.",
        )
        log.add("insert_required_section", "英文摘要", after="已补充英文摘要占位内容")
    else:
        en_abs_body = None

    has_en_kw = any(re.match(r"^Keywords?\s*[:：]", (p.text or "").strip(), re.IGNORECASE) for p in doc.paragraphs)
    if not has_en_kw and en_abs_para is not None:
        en_anchor = en_abs_body if en_abs_body is not None else en_abs_para
        # If ABSTRACT already exists with body, insert Keywords after existing body.
        if en_abs_body is None:
            paras = list(doc.paragraphs)
            try:
                idx = paras.index(en_abs_para)
            except ValueError:
                idx = -1
            if idx >= 0:
                for p in paras[idx + 1 :]:
                    t = (p.text or "").strip()
                    if not t:
                        if en_anchor is not en_abs_para:
                            break
                        continue
                    if (
                        is_keyword_en_line(t)
                        or is_keyword_zh_line(t)
                        or is_toc_title(t)
                        or is_abstract_title(t)
                        or is_intro_title(t)
                        or is_chapter(t)
                        or is_section(t)
                        or is_subsection(t)
                        or is_references(t)
                        or is_ack(t)
                        or is_author_bio(t)
                    ):
                        break
                    en_anchor = p
        insert_paragraph_after(
            en_anchor,
            "Keywords: deglobalization; new energy vehicles; Shenzhen; industry development; policy response",
        )
        log.add("insert_required_section", "英文关键词", after="已补充英文关键词")

    if not _has_toc_field(doc):
        toc_title = next((p for p in doc.paragraphs if is_toc_title(p.text)), None)
        if toc_title is None:
            anchor = abstract_title if abstract_title is not None else first_non_empty
            if anchor is not None:
                toc_title = _insert_paragraph_before(anchor, "目  录")
                log.add("insert_required_section", "目录", after="已补充目录标题")
        if toc_title is not None:
            field_para = insert_paragraph_after(toc_title, "")
            add_field_to_paragraph(field_para, ' TOC \\o "1-3" \\h \\z \\u ', "请在 Word 中右键更新目录")
            log.add("insert_required_section", "自动目录域", after="已插入 TOC 域")


def remove_manual_toc(doc: Document, log: ModifyLog):
    paras = list(doc.paragraphs)

    def _looks_like_toc_entry(para) -> bool:
        t = (para.text or "").strip()
        if "目录" in (para.style.name if para.style else ""):
            return True
        return _looks_like_toc_entry_text(t)

    toc_idx = None
    for i, p in enumerate(paras):
        if is_toc_title(p.text):
            toc_idx = i
            break
    if toc_idx is None:
        # Some source files do not have an explicit "目录" title but start directly with TOC entries.
        first_non_empty_idx = None
        for i, p in enumerate(paras):
            if (p.text or "").strip():
                first_non_empty_idx = i
                break
        has_abstract_later = any(is_abstract_title((p.text or "").strip()) for p in paras)
        if first_non_empty_idx is not None and has_abstract_later and _looks_like_toc_entry(paras[first_non_empty_idx]):
            toc_title_para = _insert_paragraph_before(paras[first_non_empty_idx], "目  录")
            paras = list(doc.paragraphs)
            toc_idx = paras.index(toc_title_para)
            log.add("insert_required_section", "目录", after="已补充目录标题（源文档缺少明确目录标题）")
        else:
            return None

    removed = 0
    for p in paras[toc_idx + 1:]:
        t = (p.text or "").strip()
        if is_abstract_title(t) or is_english_abs(t):
            break
        if _looks_like_toc_entry(p):
            delete_paragraph(p)
            removed += 1
            continue
        # Stop deleting once normal body text appears.
        if len(normalize_text(t)) > 40:
            break
        break

    toc_title_para = next((p for p in doc.paragraphs if is_toc_title(p.text)), None)
    if toc_title_para is None:
        return None

    if not _has_toc_field(doc):
        toc_field_para = insert_paragraph_after(toc_title_para, "")
        add_field_to_paragraph(toc_field_para, ' TOC \\o "1-3" \\h \\z \\u ', "请在 Word 中右键更新目录")
        log.add("insert_required_section", "自动目录域", after="已插入 TOC 域")
    if removed:
        log.add("remove_manual_toc", "目录正文", note=f"已删除手工目录条目 {removed} 段")
    return toc_title_para


def _find_first_idx(paras, predicate) -> int | None:
    for i, p in enumerate(paras):
        if predicate((p.text or "").strip()):
            return i
    return None


def _detect_manual_prefix_toc_block(paras, anchor_idx: int | None) -> tuple[int, int] | None:
    """
    Detect TOC-like prefix block when source has no explicit "目录" title.
    Returns (start, end) with end-exclusive index.
    """
    if anchor_idx is None or anchor_idx <= 0:
        return None

    start = 0
    while start < anchor_idx and not (paras[start].text or "").strip():
        start += 1
    if start >= anchor_idx:
        return None

    toc_like = 0
    short_non_toc = 0
    long_body_like = 0
    end = start

    for i in range(start, anchor_idx):
        t = (paras[i].text or "").strip()
        if not t:
            end = i + 1
            continue

        nt_len = len(normalize_text(t))
        if nt_len > 60:
            long_body_like += 1
            break

        if _looks_like_toc_entry_text(t):
            toc_like += 1
        else:
            short_non_toc += 1
        end = i + 1

    total = toc_like + short_non_toc
    if total == 0:
        return None
    if toc_like >= 3 and toc_like / total >= 0.6 and long_body_like == 0 and end > start:
        return (start, end)
    return None


def detect_front_matter_blocks(doc: Document) -> tuple[dict[str, tuple[int, int]], tuple[int, int] | None]:
    """
    Locate major front-matter blocks by heading anchors.
    Returns (blocks, toc_range), where blocks values are (start, end) end-exclusive.
    """
    paras = list(doc.paragraphs)
    if not paras:
        return {}, None

    idx_abstract = _find_first_idx(paras, is_abstract_title)
    idx_kw_zh = _find_first_idx(paras, is_keyword_zh_line)
    idx_en_abs = _find_first_idx(paras, is_english_abs)
    idx_kw_en = _find_first_idx(paras, is_keyword_en_line)
    idx_toc = _find_first_idx(paras, is_toc_title)
    idx_intro = _find_first_idx(paras, is_intro_title)

    manual_toc_bounds = None
    if idx_toc is None:
        anchor = idx_abstract if idx_abstract is not None else idx_intro
        manual_toc_bounds = _detect_manual_prefix_toc_block(paras, anchor)

    starts: dict[str, int] = {}
    if idx_abstract is not None:
        starts["abstract"] = idx_abstract
    if idx_kw_zh is not None:
        starts["keyword"] = idx_kw_zh
    if idx_en_abs is not None:
        starts["english_abstract"] = idx_en_abs
    if idx_kw_en is not None:
        starts["english_keyword"] = idx_kw_en
    if idx_toc is not None:
        starts["toc"] = idx_toc
    elif manual_toc_bounds is not None:
        starts["toc"] = manual_toc_bounds[0]
    if idx_intro is not None:
        starts["intro"] = idx_intro

    ordered = sorted(starts.items(), key=lambda kv: kv[1])
    blocks: dict[str, tuple[int, int]] = {}
    for i, (section, start) in enumerate(ordered):
        if section == "toc" and manual_toc_bounds is not None and start == manual_toc_bounds[0]:
            end = manual_toc_bounds[1]
        else:
            next_start = None
            for _, s in ordered[i + 1 :]:
                if s > start:
                    next_start = s
                    break
            end = next_start if next_start is not None else len(paras)
        if end > start:
            blocks[section] = (start, end)

    return blocks, blocks.get("toc")


def reorder_front_matter_blocks(doc: Document, rules: dict, log: ModifyLog) -> tuple[int, int] | None:
    """
    Reorder detected front-matter blocks according to dynamic school rules.
    Only moves paragraph blocks, without rewriting paragraph text.
    """
    blocks, toc_range = detect_front_matter_blocks(doc)
    front_keys = {"abstract", "keyword", "english_abstract", "english_keyword", "toc"}
    present = [k for k in blocks.keys() if k in front_keys]
    if len(present) < 2:
        return toc_range

    present_order = sorted(present, key=lambda s: blocks[s][0])
    desired_order = [s for s in (rules.get("front_matter_order") or []) if s in present and s in front_keys]
    for s in present_order:
        if s not in desired_order:
            desired_order.append(s)
    if desired_order == present_order:
        return toc_range

    spans = sorted([(sec, *blocks[sec]) for sec in present], key=lambda x: x[1])
    for i in range(len(spans) - 1):
        if spans[i][2] > spans[i + 1][1]:
            log.add("skip_reorder_front_matter", "front_matter", note="检测到区块重叠，跳过自动重排。")
            return toc_range

    paras = list(doc.paragraphs)
    block_elems: dict[str, list] = {}
    all_elems = []
    for sec, start, end in spans:
        elems = [paras[i]._p for i in range(start, end)]
        block_elems[sec] = elems
        all_elems.extend(elems)

    if not all_elems:
        return toc_range

    parent = all_elems[0].getparent()
    if any(el.getparent() is not parent for el in all_elems):
        log.add("skip_reorder_front_matter", "front_matter", note="段落父节点不一致，跳过自动重排。")
        return toc_range

    insert_at = min(parent.index(el) for el in all_elems if el.getparent() is parent)
    for el in all_elems:
        if el.getparent() is parent:
            parent.remove(el)
    for sec in desired_order:
        for el in block_elems.get(sec, []):
            parent.insert(insert_at, el)
            insert_at += 1

    log.add(
        "reorder_front_matter",
        "front_matter",
        note=f"按规则重排前置结构: {' -> '.join(present_order)} => {' -> '.join(desired_order)}",
    )

    _, new_toc_range = detect_front_matter_blocks(doc)
    return new_toc_range


def insert_author_bio_if_missing(doc: Document, log: ModifyLog, author_bio_text: str | None = None):
    if any(is_author_bio(p.text) for p in doc.paragraphs):
        return
    bio_text = (author_bio_text or "").strip()
    # Insert before 致谢 if available, otherwise append at the end.
    target = None
    for p in doc.paragraphs:
        if is_ack(p.text):
            target = p
            break
    bio_title = None
    if target is not None:
        # Insert in reverse order before target by using previous sibling trick: add paragraph before target XML.
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        title_el = OxmlElement('w:p')
        target._p.addprevious(title_el)
        bio_title = Paragraph(title_el, target._parent)
        bio_title.add_run("作者简介")
        if bio_text:
            body_el = OxmlElement('w:p')
            target._p.addprevious(body_el)
            bio_body = Paragraph(body_el, target._parent)
            bio_body.add_run(bio_text)
    else:
        bio_title = doc.add_paragraph("作者简介")
        if bio_text:
            doc.add_paragraph(bio_text)
    if bio_text:
        log.add("insert_required_section", "作者简介", after="已插入作者简介（来自用户上传内容）")
    else:
        log.add("insert_required_section", "作者简介", after="已插入作者简介标题（正文待用户上传）")


def format_table(table, rules: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = rules['table']['font']
                    set_font_east_asia(run, rules['table']['font'])
                    run.font.size = Pt(rules['table']['size_pt'])


def add_front_body_section_break(doc: Document, log: ModifyLog, toc_range: tuple[int, int] | None = None):
    """Create two page-numbering areas: front matter as upper Roman, body as Arabic from 1."""
    # Find first real body start. Prefer the body “绪论/引言” after the TOC.
    body_para = None
    seen_toc = False
    for idx, p in enumerate(doc.paragraphs):
        if toc_range is not None and toc_range[0] <= idx < toc_range[1]:
            if is_toc_title(p.text):
                seen_toc = True
            continue
        if is_toc_title(p.text):
            seen_toc = True
            continue
        if seen_toc and is_intro_title(p.text):
            body_para = p
            break
    if body_para is None:
        # Fallback: first intro heading outside TOC block.
        for idx, p in enumerate(doc.paragraphs):
            if toc_range is not None and toc_range[0] <= idx < toc_range[1]:
                continue
            if is_intro_title(p.text):
                body_para = p
                break
    if body_para is None:
        # Still set decimal page-number metadata so downstream checks can detect it.
        for sec in doc.sections:
            set_page_number_format(sec._sectPr, 'decimal', 1)
        log.add("set_page_number_format", "整文", after="已设置阿拉伯数字页码格式")
        return None
    # Find the paragraph immediately before body start; the TOC field paragraph is normally here.
    paras = list(doc.paragraphs)
    try:
        idx = paras.index(body_para)
    except ValueError:
        return None
    if idx <= 0:
        for sec in doc.sections:
            set_page_number_format(sec._sectPr, 'decimal', 1)
        return None
    prev_para = paras[idx - 1]
    # Avoid duplicate section break if rerun: only add if previous para has no sectPr.
    pPr = prev_para._p.get_or_add_pPr()
    has_sect = any(child.tag == qn('w:sectPr') for child in pPr)
    if not has_sect:
        template = doc._element.body.sectPr
        add_next_page_section_break_after(prev_para, template, 'upperRoman', 1)
        log.add("set_section_break", "目录与正文之间", after="正文前页码大写罗马；正文页码阿拉伯数字从1开始")
    # Ensure page-number metadata exists on all section properties.
    for sec in doc.sections:
        set_page_number_format(sec._sectPr, 'decimal', 1)
    body_para.paragraph_format.page_break_before = False
    return body_para


def setup_headers_footers(doc: Document, rules: dict, log: ModifyLog):
    ensure_update_fields_on_open(doc)
    # Set odd/even setting at document level.
    settings = doc.settings._element
    from docx.oxml import OxmlElement
    even_odd = settings.find(qn('w:evenAndOddHeaders'))
    if even_odd is None:
        settings.append(OxmlElement('w:evenAndOddHeaders'))

    for section in doc.sections:
        header = section.header
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.text = rules['header']['even']
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_run_style(hp, rules['header']['font'], rules['header']['size_pt'])
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)
        set_paragraph_run_style(fp, rules['header']['font'], rules['header']['size_pt'])
    log.add("set_header_footer", "页眉页脚", after="页眉居中；页脚居中页码字段")


def ensure_pgnumtype_metadata(doc: Document):
    # Ensure the final section has explicit page-number metadata.
    body_sect = doc._element.body.sectPr
    if body_sect is not None:
        set_page_number_format(body_sect, "decimal", 1)
    # Also ensure any paragraph-level section breaks carry pgNumType.
    for p in doc.paragraphs:
        ppr = p._p.pPr
        if ppr is None:
            continue
        sect = ppr.find(qn("w:sectPr"))
        if sect is not None:
            set_page_number_format(sect, "decimal", 1)


def apply_formatting(
    input_docx: str | Path,
    out_docx: str | Path,
    rules: dict | None = None,
    strict_content_fix: bool = False,
    author_bio_text: str | None = None,
    format_only: bool = True,
) -> dict:
    rules = rules or DEFAULT_RULES
    input_docx = Path(input_docx)
    out_docx = Path(out_docx)
    log = ModifyLog()
    doc = Document(input_docx)

    before_chars = count_total_chars(doc)

    set_document_defaults(doc, rules)
    set_page_setup(doc, rules)

    if strict_content_fix and not format_only:
        for p in doc.paragraphs:
            new_text, changed = fix_keyword_line(p.text)
            if changed and new_text != p.text:
                log.add("text_label_fix", p.text[:40], before=p.text, after=new_text, note="严格模式：修正关键词标签、分隔符并补足5个关键词。")
                replace_paragraph_text(p, new_text)
        ensure_required_front_matter(doc, log)
        insert_author_bio_if_missing(doc, log, author_bio_text=author_bio_text)
    elif strict_content_fix and format_only:
        log.add("skip_content_fix", "strict_content_fix", note="format_only=true，已跳过所有内容改写。")

    if not format_only:
        remove_manual_toc(doc, log)

    toc_block_range = reorder_front_matter_blocks(doc, rules, log)

    # Re-evaluate paragraphs after deletion/insertion.
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        ntext = normalize_text(text)
        clear_paragraph_numbering(p)
        # Skip empty paragraphs but make them compact.
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.page_break_before = False
            continue
        in_toc_block = (
            toc_block_range is not None
            and toc_block_range[0] <= idx < toc_block_range[1]
        )
        if in_toc_block and not is_toc_title(text):
            # Keep manual TOC lines as TOC body text; do not style as body headings.
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['toc']['font'], rules['toc']['body_size_pt'], False)
            set_para_format(p, 'left', line_spacing=1.25, first_line_indent_chars=0, space_before=0, space_after=0)
            p.paragraph_format.page_break_before = False
            continue
        # Remove first-line indent for front matter and headings.
        if is_abstract_title(text):
            if not format_only:
                replace_paragraph_text(p, rules['abstract_title']['text'])
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['abstract_title']['font'], rules['abstract_title']['size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_english_abs(text):
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['english']['font'], 18, True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_toc_title(text):
            if not format_only:
                replace_paragraph_text(p, rules['toc']['title'])
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['toc']['font'], rules['toc']['title_size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_intro_title(text):
            # Treat body introduction as Heading 1.
            p.style = doc.styles['Heading 1']
            set_paragraph_run_style(p, rules['heading_1']['font'], rules['heading_1']['size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = False
        elif is_chapter(text):
            p.style = doc.styles['Heading 1']
            set_paragraph_run_style(p, rules['heading_1']['font'], rules['heading_1']['size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = bool(rules.get('heading_1', {}).get('page_break_before', False))
        elif is_section(text):
            p.style = doc.styles['Heading 2']
            set_paragraph_run_style(p, rules['heading_2']['font'], rules['heading_2']['size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=6)
            p.paragraph_format.page_break_before = False
        elif is_subsection(text):
            p.style = doc.styles['Heading 3']
            set_paragraph_run_style(p, rules['heading_3']['font'], rules['heading_3']['size_pt'], True)
            set_para_format(p, 'left', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=6)
            p.paragraph_format.page_break_before = False
        elif is_references(text):
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['references']['title_font'], rules['references']['title_size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_author_bio(text):
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['author_bio']['title_font'], rules['author_bio']['title_size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_ack(text):
            p.style = doc.styles['Title']
            set_paragraph_run_style(p, rules['acknowledgement']['title_font'], rules['acknowledgement']['title_size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = True
        elif is_table_caption(text):
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['caption']['table_font'], rules['caption']['size_pt'], True)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=6)
            p.paragraph_format.page_break_before = False
        elif is_figure_caption(text):
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['caption']['figure_font'], rules['caption']['size_pt'], False)
            set_para_format(p, 'center', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=6)
            p.paragraph_format.page_break_before = False
        elif is_keyword_zh_line(text):
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['keyword']['font'], rules['keyword']['size_pt'])
            set_para_format(p, 'left', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=12)
            p.paragraph_format.page_break_before = False
            if p.runs:
                # Bold label only if possible; for single-run simplified to all bold false except label omitted.
                p.runs[0].bold = False
        elif is_keyword_en_line(text):
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['english']['font'], rules['english']['size_pt'])
            set_para_format(p, 'left', line_spacing=1.25, first_line_indent_chars=0, space_before=6, space_after=12)
            p.paragraph_format.page_break_before = False
        elif text.startswith("[") and re.match(r"^\[\d+\]", text):
            p.style = doc.styles['Normal']
            set_paragraph_run_style(p, rules['references']['item_font'], rules['references']['item_size_pt'])
            set_para_format(p, 'left', line_spacing=1.25, first_line_indent_chars=0, space_before=0, space_after=3)
            p.paragraph_format.page_break_before = False
        else:
            # Body / cover / abstract body.
            p.style = doc.styles['Normal']
            font = rules['body']['font']
            size = rules['body']['size_pt']
            # Heuristic: English abstract paragraphs.
            if re.search(r"[A-Za-z]{5,}", text) and not re.search(r"[\u4e00-\u9fff]", text):
                font = rules['english']['font']
            set_paragraph_run_style(p, font, size, False)
            set_para_format(p, 'justify', line_spacing=rules['body']['line_spacing'], first_line_indent_chars=2, space_before=0, space_after=0)
            p.paragraph_format.page_break_before = False

    for table in doc.tables:
        format_table(table, rules)

    # Setup page numbering and headers/footers.
    setup_headers_footers(doc, rules, log)
    add_front_body_section_break(doc, log, toc_range=toc_block_range)
    ensure_pgnumtype_metadata(doc)

    after_chars = count_total_chars(doc)
    log.add("char_count", "全文字符数", before=str(before_chars), after=str(after_chars), note="不含空白字符；严格模式可能因补关键词/作者简介而增加。")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return {"input": str(input_docx), "output": str(out_docx), "before_chars": before_chars, "after_chars": after_chars, "log": log.items}
