from __future__ import annotations

import argparse
import json
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .batch import discover_paper_files, make_case_output_dir, summarize_batch
from .calibration import calibrate_from_labels
from .engines import run_postprocess_engine
from .llm import LLMConfig, generate_suggestions
from .pipeline import run_pipeline
from .rules import extract_rules_from_text
from .scorer import save_reports, score_document


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    out.append(c.text)
    return "\n".join(out)


def read_format_text(path: str | Path) -> str:
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    converted_adjacent = path.with_name(path.stem + "_converted.docx")
    if converted_adjacent.exists():
        return read_docx_text(converted_adjacent)

    with tempfile.TemporaryDirectory() as td:
        out_dir = Path(td)
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=90,
            )
            converted = out_dir / f"{path.stem}.docx"
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass

        # Fallback: Word COM conversion on Windows.
        try:
            converted = out_dir / f"{path.stem}.docx"
            ps = rf"""
$ErrorActionPreference = "Stop"
$src = "{str(path.resolve())}"
$dst = "{str(converted.resolve())}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($src, $false, $true)
$doc.SaveAs([ref]$dst, [ref]16)
$doc.Close()
$word.Quit()
"""
            subprocess.run(
                ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=120,
            )
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass
    return ""


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Paper Format Agent CLI")
    parser.add_argument("--format-file", help="format requirement file (.doc/.docx/.txt)")
    parser.add_argument("--paper-file", help="paper file (.docx)")
    parser.add_argument("--paper-dir", help="directory of paper .docx files for batch processing")
    parser.add_argument("--paper-glob", default="*.docx", help="batch discovery glob, default: *.docx")
    parser.add_argument("--batch-fail-fast", action="store_true", help="stop batch processing after the first failure")
    parser.add_argument("--out-dir", required=True, help="output directory")
    parser.add_argument(
        "--engine",
        default="auto",
        choices=["auto", "python", "word-com", "libreoffice"],
        help="post-process engine",
    )
    parser.add_argument("--marker-dump", action="store_true", help="write paragraph type markers to marker_dump.json")
    parser.add_argument("--calibration-file", default=None, help="scoring calibration JSON")
    parser.add_argument("--strict-required-sections", action="store_true", help="enforce required sections from format file")
    parser.add_argument(
        "--allow-content-change",
        action="store_true",
        help="allow content fingerprint changes (NOT recommended for production)",
    )

    parser.add_argument("--use-llm", action="store_true", help="enable LLM suggestions (advisory only)")
    parser.add_argument("--llm-api-key", default=None)
    parser.add_argument("--llm-base-url", default="https://api.deepseek.com")
    parser.add_argument("--llm-model", default="deepseek-v4-pro")
    parser.add_argument("--llm-timeout", type=int, default=90)

    parser.add_argument("--calibrate-labels", default=None, help="manual labels JSON for calibration")
    parser.add_argument("--calibrate-out", default=None, help="output calibration JSON path")
    return parser


def run_format_job(args: argparse.Namespace, format_text: str, rules: dict, paper_file: str | Path, out_dir: Path) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    paper_file = Path(paper_file)
    (out_dir / "format_rules.json").write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

    llm_cfg = LLMConfig(
        enabled=bool(args.use_llm),
        api_key=args.llm_api_key or os.getenv("DEEPSEEK_API_KEY"),
        base_url=args.llm_base_url,
        model=args.llm_model,
        timeout_seconds=args.llm_timeout,
    )
    llm_report = generate_suggestions(paper_file, format_text, llm_cfg)
    (out_dir / "llm_suggestions.json").write_text(json.dumps(llm_report, ensure_ascii=False, indent=2), encoding="utf-8")

    marker_dump = out_dir / "marker_dump.json" if args.marker_dump else None
    output_docx = out_dir / "formatted_paper_v3.docx"
    run_result = run_pipeline(
        paper_file,
        output_docx,
        rules,
        write_marker_dump=marker_dump,
        enforce_content_guard=not bool(args.allow_content_change),
    )
    (out_dir / "modify_log.json").write_text(json.dumps(run_result.logs, ensure_ascii=False, indent=2), encoding="utf-8")

    engine_report = run_postprocess_engine(args.engine, output_docx)
    (out_dir / "engine_report.json").write_text(json.dumps(engine_report, ensure_ascii=False, indent=2), encoding="utf-8")

    report_before = score_document(
        paper_file,
        rules,
        calibration_file=args.calibration_file,
        baseline_docx=paper_file,
        enforce_required_sections=bool(args.strict_required_sections),
    )
    report_after = score_document(
        output_docx,
        rules,
        calibration_file=args.calibration_file,
        baseline_docx=paper_file,
        enforce_required_sections=bool(args.strict_required_sections),
    )

    report = report_after.copy()
    report["score_before"] = round(report_before["score"], 1)
    report["score_after"] = round(report_after["score"], 1)
    report["score_improvement"] = round(report_after["score"] - report_before["score"], 1)
    report["chars_no_space_before"] = report_before["chars_no_space"]
    report["chars_no_space_after"] = report_after["chars_no_space"]
    report["llm_used"] = bool(llm_report.get("used"))
    report["llm_warnings"] = llm_report.get("warnings", [])
    report["engine_report"] = engine_report
    report["removed_numpr_count"] = run_result.removed_numpr_count
    report["classification_confidence"] = run_result.classification_confidence
    report["content_fingerprint_before"] = run_result.content_fingerprint_before
    report["content_fingerprint_after"] = run_result.content_fingerprint_after
    report["content_changed"] = bool(run_result.content_changed)
    report["content_guard_enforced"] = not bool(args.allow_content_change)
    save_reports(report, out_dir / "format_report.json", out_dir / "format_report.html")

    return {
        "paper_file": str(paper_file),
        "out_dir": str(out_dir),
        "output": str(output_docx),
        "score_before": report["score_before"],
        "score_after": report["score_after"],
        "score_improvement": report["score_improvement"],
        "raw_quality_score": report.get("raw_quality_score"),
        "chars_no_space_before": report["chars_no_space_before"],
        "chars_no_space_after": report["chars_no_space_after"],
        "removed_numpr_count": run_result.removed_numpr_count,
        "content_changed": bool(run_result.content_changed),
        "content_guard_enforced": not bool(args.allow_content_change),
        "engine": engine_report.get("engine", args.engine),
        "engine_success": bool(engine_report.get("success")),
        "llm_used": bool(llm_report.get("used")),
    }


def run_batch(args: argparse.Namespace, format_text: str, rules: dict, out_dir: Path) -> int:
    if not args.paper_dir:
        raise ValueError("batch mode requires --paper-dir")

    paper_dir = Path(args.paper_dir)
    papers = discover_paper_files(paper_dir, args.paper_glob)
    if not papers:
        summary = summarize_batch([])
        summary["error"] = f"no .docx papers found in {paper_dir}"
        (out_dir / "batch_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 2

    results: list[dict] = []
    for paper in papers:
        case_out_dir = make_case_output_dir(out_dir, paper, paper_dir)
        try:
            result = run_format_job(args, format_text, rules, paper, case_out_dir)
            result["ok"] = not bool(result.get("content_changed")) and bool(result.get("engine_success"))
        except Exception as exc:
            result = {
                "paper_file": str(paper),
                "out_dir": str(case_out_dir),
                "ok": False,
                "error": str(exc),
            }
            if args.batch_fail_fast:
                results.append(result)
                break
        results.append(result)

    summary = summarize_batch(results)
    (out_dir / "batch_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if summary["failed"] == 0 else 1


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if args.calibrate_labels:
        out_file = Path(args.calibrate_out) if args.calibrate_out else (out_dir / "scoring_calibration.json")
        rules = extract_rules_from_text(read_format_text(args.format_file)) if args.format_file else extract_rules_from_text("")
        result = calibrate_from_labels(args.calibrate_labels, out_file, rules=rules)
        print(json.dumps({"calibration_file": str(out_file), **result}, ensure_ascii=False, indent=2))
        return 0

    if args.paper_file and args.paper_dir:
        raise ValueError("use either --paper-file or --paper-dir, not both")
    if not args.format_file:
        raise ValueError("formatting mode requires --format-file")

    format_text = read_format_text(args.format_file)
    rules = extract_rules_from_text(format_text)

    if args.paper_dir:
        return run_batch(args, format_text, rules, out_dir)

    if not args.paper_file:
        raise ValueError("normal formatting mode requires --paper-file")

    result = run_format_job(args, format_text, rules, args.paper_file, out_dir)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
