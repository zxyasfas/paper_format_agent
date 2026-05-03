from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib import error, request

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .ooxml import insert_paragraph_after


@dataclass
class LLMConfig:
    enabled: bool = False
    api_key: str | None = None
    base_url: str = "https://api.deepseek.com"
    model: str = "deepseek-v4-pro"
    timeout_seconds: int = 90
    apply_content_fixes: bool = False
    max_format_chars: int = 4000
    max_paper_chars: int = 12000


def _norm(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _looks_like_heading(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    if len(t) <= 30:
        if re.match(r"^[一二三四五六七八九十]+、", t):
            return True
        if re.match(r"^第[一二三四五六七八九十0-9]+章", t):
            return True
        if re.match(r"^\d+(\.\d+){0,2}\s*", t):
            return True
        if any(k in t for k in ("绪论", "引言", "结论", "结语", "参考文献", "致谢")):
            return True
    return False


def _collect_heading_candidates(doc: Document) -> list[str]:
    seen: set[str] = set()
    items: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading") or _looks_like_heading(t):
            n = _norm(t)
            if n not in seen:
                seen.add(n)
                items.append(t)
    return items[:80]


def _extract_abstract_text(doc: Document) -> str:
    paras = doc.paragraphs
    abstract_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if re.fullmatch(r"摘\s*要", t):
            abstract_idx = i
            break

    if abstract_idx is not None:
        chunks: list[str] = []
        for p in paras[abstract_idx + 1 :]:
            t = (p.text or "").strip()
            if not t:
                if chunks:
                    break
                continue
            if _looks_like_heading(t):
                break
            chunks.append(t)
            if sum(len(x) for x in chunks) > 1600:
                break
        if chunks:
            return "\n".join(chunks)

    for p in paras:
        t = (p.text or "").strip()
        if len(t) >= 80 and re.search(r"[\u4e00-\u9fff]", t):
            return t[:1600]
    return ""


def _find_first_paragraph(doc: Document, pattern: str, flags: int = 0):
    cre = re.compile(pattern, flags)
    for p in doc.paragraphs:
        if cre.match((p.text or "").strip()):
            return p
    return None


def _insert_before(target_paragraph, text: str):
    new_p = OxmlElement("w:p")
    target_paragraph._p.addprevious(new_p)
    para = Paragraph(new_p, target_paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _find_by_exact_text(doc: Document, text: str):
    nt = _norm(text)
    for p in doc.paragraphs:
        if _norm((p.text or "").strip()) == nt:
            return p
    return None


def _has_intro(doc: Document) -> bool:
    for p in doc.paragraphs:
        t = _norm((p.text or "").strip())
        if "绪论" in t or "引言" in t:
            return True
    return False


def _has_conclusion(doc: Document) -> bool:
    for p in doc.paragraphs:
        t = _norm((p.text or "").strip())
        if "结论" in t or "结语" in t:
            return True
    return False


def _has_keywords_zh(doc: Document) -> bool:
    return _find_first_paragraph(doc, r"^(关键词|关键字)\s*[:：]") is not None


def _has_abstract_en(doc: Document) -> bool:
    return _find_first_paragraph(doc, r"^ABSTRACT$", flags=re.IGNORECASE) is not None


def _has_keywords_en(doc: Document) -> bool:
    return _find_first_paragraph(doc, r"^Keywords?\s*[:：]", flags=re.IGNORECASE) is not None


def _build_prompt(format_text: str, abstract_text: str, headings: list[str], missing: dict[str, bool]) -> str:
    headings_block = "\n".join(f"- {h}" for h in headings[:80])
    return (
        "You are helping a thesis formatting system.\n"
        "Return strict JSON only with keys:\n"
        "{\n"
        '  "intro_heading_candidates": [string],\n'
        '  "conclusion_heading_candidates": [string],\n'
        '  "keywords_zh": [string],\n'
        '  "english_abstract": "string",\n'
        '  "english_keywords": [string]\n'
        "}\n"
        "Rules:\n"
        "- Pick candidates from the given heading list only.\n"
        "- keywords_zh: 5-8 Chinese keywords, noun phrases.\n"
        "- english_abstract: translate/rewrite from Chinese abstract, 120-220 words, academic tone.\n"
        "- english_keywords: 5-8 keywords matching Chinese ones.\n"
        "- If uncertain, return empty list/string for that field.\n\n"
        f"Missing fields now: {json.dumps(missing, ensure_ascii=False)}\n\n"
        f"Format requirement excerpt:\n{format_text}\n\n"
        f"Chinese abstract/source excerpt:\n{abstract_text}\n\n"
        f"Heading candidates:\n{headings_block}\n"
    )


def _call_openai_compatible(cfg: LLMConfig, prompt: str) -> str:
    if not cfg.api_key:
        raise RuntimeError("LLM API key is missing.")

    body = {
        "model": cfg.model,
        "temperature": 0.2,
        "reasoning_effort": "high",
        "thinking": {"type": "enabled"},
        "response_format": {"type": "json_object"},
        "messages": [
            {
                "role": "system",
                "content": (
                    "You output strict JSON only. No markdown. "
                    "Do not invent heading strings not present in candidates."
                ),
            },
            {"role": "user", "content": prompt},
        ],
    }
    payload = json.dumps(body).encode("utf-8")
    endpoint = cfg.base_url.rstrip("/") + "/chat/completions"
    req = request.Request(
        endpoint,
        data=payload,
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {cfg.api_key}",
        },
    )
    try:
        with request.urlopen(req, timeout=cfg.timeout_seconds) as resp:
            raw = resp.read().decode("utf-8")
    except error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"LLM HTTP {e.code}: {detail[:500]}") from e
    except Exception as e:
        raise RuntimeError(f"LLM request failed: {e}") from e

    data = json.loads(raw)
    content = data["choices"][0]["message"]["content"]
    if isinstance(content, list):
        chunks = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                chunks.append(item.get("text", ""))
        content = "".join(chunks)
    if not isinstance(content, str):
        raise RuntimeError("Unexpected LLM response content format.")
    return content.strip()


def _safe_parse_json(text: str) -> dict[str, Any]:
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not m:
            return {}
        try:
            return json.loads(m.group(0))
        except Exception:
            return {}


def _sanitize_list(v: Any, limit: int = 8) -> list[str]:
    if not isinstance(v, list):
        return []
    out: list[str] = []
    for item in v:
        if isinstance(item, str):
            s = item.strip()
            if s:
                out.append(s)
        if len(out) >= limit:
            break
    return out


def enhance_docx_with_llm(
    paper_file: str | Path,
    format_text: str,
    out_dir: str | Path,
    cfg: LLMConfig,
) -> dict[str, Any]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    paper_file = Path(paper_file)
    report: dict[str, Any] = {
        "enabled": cfg.enabled,
        "used": False,
        "model": cfg.model,
        "base_url": cfg.base_url,
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "warnings": [],
        "actions": [],
        "response": {},
    }

    if not cfg.enabled:
        report["input_docx"] = str(paper_file)
        report["output_docx"] = str(paper_file)
        return report

    doc = Document(paper_file)
    headings = _collect_heading_candidates(doc)
    abstract_text = _extract_abstract_text(doc)[: cfg.max_paper_chars]
    format_excerpt = (format_text or "")[: cfg.max_format_chars]
    missing = {
        "keywords_zh": not _has_keywords_zh(doc),
        "english_abstract": not _has_abstract_en(doc),
        "english_keywords": not _has_keywords_en(doc),
        "intro_heading": not _has_intro(doc),
        "conclusion_heading": not _has_conclusion(doc),
    }
    report["input_docx"] = str(paper_file)
    report["missing_before"] = missing

    prompt = _build_prompt(format_excerpt, abstract_text, headings, missing)
    try:
        raw_content = _call_openai_compatible(cfg, prompt)
        parsed = _safe_parse_json(raw_content)
        report["used"] = True
        report["response_raw"] = raw_content
        report["response"] = parsed
    except Exception as e:
        msg = str(e)
        report["warnings"].append(msg)
        # Reliability fallback: v4-pro timeout -> v4-flash
        if cfg.model == "deepseek-v4-pro" and "timed out" in msg.lower():
            fallback_cfg = LLMConfig(
                enabled=cfg.enabled,
                api_key=cfg.api_key,
                base_url=cfg.base_url,
                model="deepseek-v4-flash",
                timeout_seconds=cfg.timeout_seconds,
                apply_content_fixes=cfg.apply_content_fixes,
                max_format_chars=cfg.max_format_chars,
                max_paper_chars=cfg.max_paper_chars,
            )
            try:
                raw_content = _call_openai_compatible(fallback_cfg, prompt)
                parsed = _safe_parse_json(raw_content)
                report["used"] = True
                report["model"] = fallback_cfg.model
                report["response_raw"] = raw_content
                report["response"] = parsed
                report["warnings"].append("Fallback to deepseek-v4-flash after v4-pro timeout.")
            except Exception as e2:
                report["warnings"].append(str(e2))
                report["output_docx"] = str(paper_file)
                return report
        else:
            report["output_docx"] = str(paper_file)
            return report

    intro_candidates = _sanitize_list(report["response"].get("intro_heading_candidates"))
    conclusion_candidates = _sanitize_list(report["response"].get("conclusion_heading_candidates"))
    keywords_zh = _sanitize_list(report["response"].get("keywords_zh"))
    en_keywords = _sanitize_list(report["response"].get("english_keywords"))
    en_abs = report["response"].get("english_abstract")
    if not isinstance(en_abs, str):
        en_abs = ""
    en_abs = en_abs.strip()

    report["suggestions"] = {
        "intro_heading_candidates": intro_candidates,
        "conclusion_heading_candidates": conclusion_candidates,
        "keywords_zh": keywords_zh,
        "english_abstract": en_abs,
        "english_keywords": en_keywords,
    }
    if not cfg.apply_content_fixes:
        # Conservative default: do not mutate thesis content automatically.
        report["warnings"].append("LLM content edits disabled (suggestions-only mode).")
        report["output_docx"] = str(paper_file)
        return report

    # Semantic weak-structure fixes: add canonical labels when missing.
    if not _has_intro(doc) and intro_candidates:
        target = _find_by_exact_text(doc, intro_candidates[0])
        if target is not None:
            _insert_before(target, "引言")
            report["actions"].append({"action": "insert_semantic_heading", "target": "引言", "anchor": intro_candidates[0]})
    if not _has_conclusion(doc) and conclusion_candidates:
        target = _find_by_exact_text(doc, conclusion_candidates[0])
        if target is not None:
            _insert_before(target, "结论")
            report["actions"].append({"action": "insert_semantic_heading", "target": "结论", "anchor": conclusion_candidates[0]})

    # Front-matter semantic fixes: keywords + English abstract + English keywords.
    has_zh_kw = _has_keywords_zh(doc)
    has_en_abs = _has_abstract_en(doc)
    has_en_kw = _has_keywords_en(doc)

    anchor = _find_first_paragraph(doc, r"^摘\s*要$")
    if anchor is None:
        anchor = next((p for p in doc.paragraphs if (p.text or "").strip()), None)
    if anchor is None:
        report["warnings"].append("Document has no non-empty paragraph.")
        report["output_docx"] = str(paper_file)
        return report

    # Move anchor to end of abstract body if possible.
    paras = list(doc.paragraphs)
    try:
        idx = paras.index(anchor)
        for p in paras[idx + 1 :]:
            t = (p.text or "").strip()
            if not t:
                if p is not anchor:
                    break
                continue
            if _looks_like_heading(t):
                break
            anchor = p
    except ValueError:
        pass

    if not has_zh_kw and keywords_zh:
        kw_line = "关键词：" + "；".join(keywords_zh[:8])
        anchor = insert_paragraph_after(anchor, kw_line)
        report["actions"].append({"action": "insert_keywords_zh", "value": kw_line})

    if not has_en_abs and en_abs:
        anchor = insert_paragraph_after(anchor, "ABSTRACT")
        anchor = insert_paragraph_after(anchor, en_abs)
        report["actions"].append({"action": "insert_english_abstract", "chars": len(en_abs)})

    if not has_en_kw and en_keywords:
        kw_en_line = "Keywords: " + "; ".join(en_keywords[:8])
        anchor = insert_paragraph_after(anchor, kw_en_line)
        report["actions"].append({"action": "insert_keywords_en", "value": kw_en_line})

    enhanced_docx = out_dir / "paper_llm_enhanced.docx"
    doc.save(enhanced_docx)
    report["output_docx"] = str(enhanced_docx)
    return report


def config_from_env_or_args(
    *,
    use_llm: bool,
    llm_api_key: str | None,
    llm_base_url: str | None,
    llm_model: str | None,
    llm_timeout: int,
    apply_llm_content_fixes: bool = False,
) -> LLMConfig:
    api_key = (
        llm_api_key
        or os.getenv("DEEPSEEK_API_KEY")
        or os.getenv("PAPER_FORMAT_LLM_API_KEY")
        or os.getenv("OPENAI_API_KEY")
    )
    base_url = (
        llm_base_url
        or os.getenv("DEEPSEEK_BASE_URL")
        or os.getenv("PAPER_FORMAT_LLM_BASE_URL")
        or os.getenv("OPENAI_BASE_URL")
        or "https://api.deepseek.com"
    )
    model = llm_model or os.getenv("PAPER_FORMAT_LLM_MODEL") or os.getenv("DEEPSEEK_MODEL") or "deepseek-v4-pro"
    return LLMConfig(
        enabled=use_llm,
        api_key=api_key,
        base_url=base_url,
        model=model,
        timeout_seconds=llm_timeout,
        apply_content_fixes=apply_llm_content_fixes,
    )
