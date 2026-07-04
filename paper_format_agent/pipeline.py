from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .ooxml import set_font_east_asia
from .quality import build_content_fingerprint


MARK_PREFIX = "PFA3_MARK_"

T_EMPTY = "EMPTY"
T_BODY = "BODY"
T_TOC_TITLE = "TOC_TITLE"
T_TOC_ENTRY = "TOC_ENTRY"
T_ABS_ZH_TITLE = "ABS_ZH_TITLE"
T_ABS_ZH_BODY = "ABS_ZH_BODY"
T_KW_ZH = "KW_ZH"
T_ABS_EN_TITLE = "ABS_EN_TITLE"
T_ABS_EN_BODY = "ABS_EN_BODY"
T_KW_EN = "KW_EN"
T_H1 = "H1"
T_H2 = "H2"
T_H3 = "H3"
T_REF_TITLE = "REF_TITLE"
T_REF_ENTRY = "REF_ENTRY"
T_ACK_TITLE = "ACK_TITLE"
T_FIG_CAPTION = "FIG_CAPTION"
T_TABLE_CAPTION = "TABLE_CAPTION"


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def is_abstract_title(text: str) -> bool:
    # Strip whitespace, punctuation wrappers, and optional "\u4e2d\u6587" prefix.
    t = normalize_text(text)
    t = re.sub(r"^[\uff08\uff09\[\]\u3010\u3011\uff3b\uff3d\u3014\u3015\u3016\u3017\uff1a:\s]+|[\uff08\uff09\[\]\u3010\u3011\uff3b\uff3d\u3014\u3015\u3016\u3017\uff1a:\s]+$", "", t)
    t = re.sub(r"^\u4e2d\u6587", "", t)  # strip \u4e2d\u6587 prefix
    return t == "\u6458\u8981"  # \u6458\u8981


def is_english_abstract_title(text: str) -> bool:
    return (text or "").strip().upper() == "ABSTRACT"


def is_keyword_zh(text: str) -> bool:
    # Chinese keywords line. Strip whitespace and punctuation wrappers, then
    # match the common "关键词" / "关键字" label with an optional "中文" prefix
    # and optional trailing colon.
    t = normalize_text(text)
    t = re.sub(r"^[（）\[\]【】［］〔〕〖〗：:\s]+|[（）\[\]【】［］〔〕〖〗：:\s]+$", "", t)
    t = re.sub(r"^中文", "", t)  # strip 中文 prefix
    return bool(re.match(r"^(关键词|关键字)(?:[:：]|$)", t))


def is_keyword_en(text: str) -> bool:
    return bool(re.match(r"^Keywords?\s*[:\uff1a]", (text or "").strip(), flags=re.IGNORECASE))


def is_toc_title(text: str) -> bool:
    return normalize_text(text) in {"\u76ee\u5f55", "\u76ee\u6b21", "\u76ee\u5f55\u9875"}


def is_references(text: str) -> bool:
    nt = normalize_text(text)
    return ("\u53c2\u8003\u6587\u732e" in nt) or ("references" in nt.lower())


def is_ack(text: str) -> bool:
    nt = normalize_text(text)
    return ("\u81f4\u8c22" in nt) or ("\u81f4\u8b1d" in nt) or ("acknowledg" in nt.lower())


def is_intro_like(text: str) -> bool:
    nt = normalize_text(text)
    return ("\u7eea\u8bba" in nt) or ("\u5f15\u8a00" in nt) or (nt.lower() == "introduction")


def looks_like_reference_entry(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    if re.match(r"^\[\d+\]", t):
        return True
    if re.match(r"^\d+\s*[\.\u3001\uff0e)]", t):
        return True
    if "DOI" in t.upper():
        return True
    if "[J]" in t or "[M]" in t or "[D]" in t:
        return True
    return False


def is_chapter(text: str) -> bool:
    s = (text or "").strip()
    return bool(
        re.match(
            r"^(\u7b2c[\u4e00-\u9fff0-9]+[\u7ae0\u8282]|[\u4e00-\u9fff]{1,4}[\u3001.\uff0e]|[IVXLC]+\s*[.\u3001])\s*",
            s,
        )
    )


def is_section(text: str) -> bool:
    s = (text or "").strip()
    return bool(re.match(r"^(?:\d+\.\d+(?!\.)\s*|[A-Z]\.(?:\s+|$))", s))


def is_subsection(text: str) -> bool:
    s = (text or "").strip()
    return bool(re.match(r"^\d+\.\d+\.\d+\s*", s))


def looks_like_toc_entry(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    nt = normalize_text(t)
    if re.search(r"[\u00b7.\u3002\u2026]{3,}\s*\d+\s*$", t):
        return True
    if bool(re.match(r"^[\u4e00-\u9fff]{1,6}[\u3001.\uff0e]", t)):
        return True
    if bool(re.match(r"^[\uff08(][\u4e00-\u9fff]{1,4}[\uff09)]", t)):
        return True
    if bool(re.match(r"^\d+(\.\d+)*\s+", t)):
        return True
    if bool(re.match(r"^\d+\.(\d+\.)*\d*[^0-9\s]", t)):
        return True
    if bool(re.match(r"^\u7b2c[\u4e00-\u9fff0-9]+[\u7ae0\u8282]", t)):
        return True
    if len(nt) <= 40 and (is_chapter(t) or is_section(t) or is_subsection(t) or is_intro_like(t)):
        return True
    return False


def is_figure_caption(text: str) -> bool:
    t = (text or "").strip()
    if len(normalize_text(t)) > 90:
        return False
    return bool(
        re.match(
            r"^(\u56fe|Fig\.?)\s*[0-9\u4e00-\u9fff.\-\uFF0D\u2014]+(?:\s|[:\uff1a\uFF0E.\u3001]|$)",
            t,
            flags=re.IGNORECASE,
        )
    )


def is_table_caption(text: str) -> bool:
    t = (text or "").strip()
    if len(normalize_text(t)) > 90:
        return False
    return bool(
        re.match(
            r"^(\u8868|Table)\s*[0-9\u4e00-\u9fffIVXLCDM.\-\uFF0D\u2014]+(?:\s|[:\uff1a\uFF0E.\u3001]|$)",
            t,
            flags=re.IGNORECASE,
        )
    )


def detect_toc_range_after_title(paras: list, title_idx: int) -> tuple[int, int] | None:
    """Detect TOC entry block following a TOC title."""
    start: int | None = None
    last: int | None = None
    non_empty = 0
    toc_like = 0
    empty_run = 0

    for j in range(title_idx + 1, len(paras)):
        t = (paras[j].text or "").strip()
        if not t:
            empty_run += 1
            if start is not None and toc_like >= 5 and empty_run >= 3:
                break
            continue

        empty_run = 0
        nt = normalize_text(t)
        nt_len = len(nt)
        sentence_marks = len(re.findall(r"[\u3002\uff01\uff1f!?\uff1b;]", nt))
        strong_body = nt_len >= 80 or (nt_len >= 56 and sentence_marks >= 2)

        if is_toc_title(t) or is_abstract_title(t) or is_english_abstract_title(t) or is_references(t) or is_ack(t):
            break

        chapter_like = is_chapter(t) or is_section(t) or is_subsection(t) or is_intro_like(t)
        entry_like = looks_like_toc_entry(t) or chapter_like or nt_len <= 42

        if start is None:
            if not entry_like:
                return None
            start = j
            last = j
            non_empty += 1
            toc_like += 1
            continue

        if strong_body and toc_like >= 5:
            break

        if entry_like or (nt_len <= 60 and sentence_marks <= 1 and not strong_body):
            non_empty += 1
            if entry_like:
                toc_like += 1
            last = j
            continue

        if toc_like >= 5:
            break
        return None

    if start is None or last is None or non_empty < 4:
        return None
    ratio = toc_like / non_empty if non_empty else 0.0
    if ratio < 0.68:
        return None
    return (start, last + 1)


def clear_paragraph_numbering(paragraph) -> bool:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)
        return True
    return False


def is_in_table_cell(paragraph) -> bool:
    try:
        return bool(paragraph._p.xpath("./ancestor::w:tc"))
    except Exception:
        return False


def collapse_soft_breaks(text: str, joiner: str = "") -> str:
    if not text:
        return text
    if ("\n" not in text) and ("\r" not in text):
        return text
    parts = [x.strip() for x in re.split(r"[\r\n]+", text) if x.strip()]
    if not parts:
        return ""
    return joiner.join(parts)


def ensure_style(doc: Document, name: str, font: str, size_pt: float, bold: bool = False, alignment: str | None = None):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    if alignment == "center":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "left":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == "justify":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return style


def set_run_style(paragraph, font: str, size_pt: float, bold: bool | None = None):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = font
        set_font_east_asia(run, font)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_para_format(paragraph, *, align: str, line_spacing: float, indent_chars: int, space_before: int, space_after: int):
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(12 * indent_chars)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def _clear_paragraph(paragraph):
    paragraph.text = ""


def _add_page_number_field(paragraph, font: str, size_pt: float):
    run = paragraph.add_run()
    run.font.name = font
    set_font_east_asia(run, font)
    run.font.size = Pt(size_pt)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def apply_header_footer(doc: Document, rules: dict) -> dict[str, int]:
    header_cfg = rules.get("header", {})
    footer_cfg = rules.get("footer", {})
    header_text = str(header_cfg.get("text", "") or "").strip()
    header_font = header_cfg.get("font", rules["body"]["font"])
    header_size = float(header_cfg.get("size_pt", 9))
    footer_font = footer_cfg.get("font", header_font)
    footer_size = float(footer_cfg.get("size_pt", header_size))
    add_page_number = bool(footer_cfg.get("page_number", True))

    headers_written = 0
    footers_written = 0
    for section in doc.sections:
        if header_text:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            _clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(header_text)
            run.font.name = header_font
            set_font_east_asia(run, header_font)
            run.font.size = Pt(header_size)
            headers_written += 1

        if add_page_number:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            _clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_page_number_field(paragraph, footer_font, footer_size)
            footers_written += 1

    return {"headers_written": headers_written, "footers_written": footers_written}


def setup_document_base(doc: Document, rules: dict) -> dict[str, int]:
    normal = doc.styles["Normal"]
    normal.font.name = rules["body"]["font"]
    normal.font.size = Pt(rules["body"]["size_pt"])
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), rules["body"]["font"])
    rfonts.set(qn("w:ascii"), rules["body"]["font"])
    rfonts.set(qn("w:hAnsi"), rules["body"]["font"])

    ensure_style(doc, "Title", rules["toc"]["font"], rules["toc"]["title_size_pt"], True, "center")
    ensure_style(
        doc,
        "Heading 1",
        rules["heading_1"]["font"],
        rules["heading_1"]["size_pt"],
        True,
        rules["heading_1"].get("align", "center"),
    )
    ensure_style(
        doc,
        "Heading 2",
        rules["heading_2"]["font"],
        rules["heading_2"]["size_pt"],
        True,
        rules["heading_2"].get("align", "left"),
    )
    ensure_style(
        doc,
        "Heading 3",
        rules["heading_3"]["font"],
        rules["heading_3"]["size_pt"],
        True,
        rules["heading_3"].get("align", "left"),
    )

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(rules["margins_cm"]["top"])
        section.bottom_margin = Cm(rules["margins_cm"]["bottom"])
        section.left_margin = Cm(rules["margins_cm"]["left"])
        section.right_margin = Cm(rules["margins_cm"]["right"])
    return apply_header_footer(doc, rules)


@dataclass
class Classification:
    types: list[str]
    prefix_toc_range: tuple[int, int] | None = None
    confidence: float = 0.0
    notes: list[str] = field(default_factory=list)


def detect_prefix_toc_range(doc: Document) -> tuple[int, int] | None:
    paras = list(doc.paragraphs)
    idx_abs = None
    for i, p in enumerate(paras):
        if is_abstract_title((p.text or "").strip()):
            idx_abs = i
            break
    if idx_abs is None or idx_abs < 8:
        return None

    start = 0
    while start < idx_abs and not (paras[start].text or "").strip():
        start += 1
    if start >= idx_abs:
        return None

    total = 0
    toc_like = 0
    short_like = 0
    body_like = 0
    for p in paras[start:idx_abs]:
        t = (p.text or "").strip()
        if not t:
            continue
        nt = normalize_text(t)
        total += 1
        if len(nt) <= 42:
            short_like += 1
        sentence_marks = len(re.findall(r"[。！？!?；;：:]", nt))
        if len(nt) >= 80 or (len(nt) >= 56 and sentence_marks >= 2):
            body_like += 1
        if looks_like_toc_entry(t):
            toc_like += 1
    if total < 8:
        return None
    toc_ratio = toc_like / total
    short_ratio = short_like / total
    if toc_ratio >= 0.72 and short_ratio >= 0.78 and body_like <= max(1, total // 12):
        return (start, idx_abs)
    return None


def classify_document(doc: Document) -> Classification:
    paras = list(doc.paragraphs)
    types = [T_EMPTY for _ in paras]
    notes: list[str] = []

    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            types[i] = T_EMPTY
            continue
        types[i] = T_BODY

    prefix_toc_range = detect_prefix_toc_range(doc)
    if prefix_toc_range:
        s, e = prefix_toc_range
        for i in range(s, e):
            if (paras[i].text or "").strip():
                types[i] = T_TOC_ENTRY
        notes.append(f"prefix_toc_detected={prefix_toc_range}")

    # Strong anchors
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        if prefix_toc_range is not None and prefix_toc_range[0] <= i < prefix_toc_range[1]:
            if is_toc_title(t):
                types[i] = T_TOC_TITLE
            else:
                types[i] = T_TOC_ENTRY
            continue
        if is_toc_title(t):
            types[i] = T_TOC_TITLE
        elif is_abstract_title(t):
            types[i] = T_ABS_ZH_TITLE
        elif is_english_abstract_title(t):
            types[i] = T_ABS_EN_TITLE
        elif is_keyword_zh(t):
            types[i] = T_KW_ZH
        elif is_keyword_en(t):
            types[i] = T_KW_EN
        elif is_references(t):
            types[i] = T_REF_TITLE
        elif is_ack(t):
            types[i] = T_ACK_TITLE
        elif is_figure_caption(t):
            types[i] = T_FIG_CAPTION
        elif is_table_caption(t):
            types[i] = T_TABLE_CAPTION

    # Mark TOC entries right after each TOC title.
    toc_ranges: list[tuple[int, int]] = []
    for i, tp in enumerate(types):
        if tp != T_TOC_TITLE:
            continue
        rng = detect_toc_range_after_title(paras, i)
        if rng is None:
            continue
        s, e = rng
        toc_ranges.append(rng)
        for j in range(s, e):
            if (paras[j].text or "").strip():
                types[j] = T_TOC_ENTRY
    if toc_ranges:
        notes.append(f"toc_ranges_after_title={toc_ranges}")

    # Heading class for non-front-matter lines.
    for i, p in enumerate(paras):
        if types[i] in {
            T_TOC_TITLE,
            T_TOC_ENTRY,
            T_ABS_ZH_TITLE,
            T_ABS_EN_TITLE,
            T_KW_ZH,
            T_KW_EN,
            T_REF_TITLE,
            T_ACK_TITLE,
            T_FIG_CAPTION,
            T_TABLE_CAPTION,
            T_EMPTY,
        }:
            continue
        t = (p.text or "").strip()
        if is_chapter(t) or is_intro_like(t):
            types[i] = T_H1
        elif is_subsection(t):
            types[i] = T_H3
        elif is_section(t):
            types[i] = T_H2

    # Mark abstract bodies via boundary scan.
    def mark_body_after(title_type: str, body_type: str, break_types: set[str]):
        idx = next((i for i, tp in enumerate(types) if tp == title_type), None)
        if idx is None:
            return
        for j in range(idx + 1, len(types)):
            if types[j] in break_types:
                break
            t = (paras[j].text or "").strip()
            if not t:
                if body_type == T_ABS_ZH_BODY or body_type == T_ABS_EN_BODY:
                    if j > idx + 1:
                        break
                    continue
                continue
            if types[j] in {T_BODY, T_H1, T_H2, T_H3}:
                # Abstract body can include plain lines; stop at heading-like boundaries.
                if types[j] in {T_H1, T_H2, T_H3}:
                    break
                types[j] = body_type

    mark_body_after(
        T_ABS_ZH_TITLE,
        T_ABS_ZH_BODY,
        {T_KW_ZH, T_ABS_EN_TITLE, T_TOC_TITLE, T_REF_TITLE, T_ACK_TITLE, T_H1},
    )
    mark_body_after(
        T_ABS_EN_TITLE,
        T_ABS_EN_BODY,
        {T_KW_EN, T_TOC_TITLE, T_REF_TITLE, T_ACK_TITLE, T_H1},
    )

    # Mark references block entries to avoid justify/distributed spacing artifacts.
    ref_title_indices = [i for i, tp in enumerate(types) if tp == T_REF_TITLE]
    idx_ref_title = max(ref_title_indices) if ref_title_indices else None
    if idx_ref_title is not None:
        for j in range(idx_ref_title + 1, len(types)):
            t = (paras[j].text or "").strip()
            if not t:
                continue
            if is_ack(t) or types[j] == T_ACK_TITLE:
                break
            if types[j] in {T_ABS_ZH_TITLE, T_ABS_EN_TITLE, T_TOC_TITLE, T_REF_TITLE}:
                break
            if types[j] in {T_H1, T_H2} and not looks_like_reference_entry(t):
                break
            if looks_like_reference_entry(t) or types[j] in {T_BODY, T_H3}:
                types[j] = T_REF_ENTRY

    confidence = 0.0
    if next((i for i, tp in enumerate(types) if tp == T_ABS_ZH_TITLE), None) is not None:
        confidence += 0.35
    if prefix_toc_range is not None:
        confidence += 0.45
    elif next((i for i, tp in enumerate(types) if tp == T_TOC_TITLE), None) is not None:
        confidence += 0.30
    if next((i for i, tp in enumerate(types) if tp == T_H1), None) is not None:
        confidence += 0.25
    confidence = min(1.0, confidence)

    return Classification(types=types, prefix_toc_range=prefix_toc_range, confidence=confidence, notes=notes)


def reorder_by_types(doc: Document, cls: Classification) -> tuple[list[dict[str, Any]], tuple[int, int] | None]:
    """
    Move detected prefix TOC block after abstract/keywords/english-abstract block.
    This is only applied when confidence is high enough.
    """
    logs: list[dict[str, Any]] = []
    if cls.prefix_toc_range is None or cls.confidence < 0.75:
        if cls.prefix_toc_range is not None:
            logs.append({"action": "skip_reorder_low_confidence", "confidence": cls.confidence})
        return logs, None

    s, e = cls.prefix_toc_range
    paras = list(doc.paragraphs)
    if e <= s or e > len(paras):
        return logs, None

    # Find anchor end in front matter (zh abstract -> zh kw -> en abstract -> en kw).
    anchor_candidates = []
    for i, tp in enumerate(cls.types):
        if tp in {T_ABS_ZH_TITLE, T_ABS_ZH_BODY, T_KW_ZH, T_ABS_EN_TITLE, T_ABS_EN_BODY, T_KW_EN}:
            anchor_candidates.append(i)
    if not anchor_candidates:
        return logs, None

    anchor_idx = max(anchor_candidates)
    if anchor_idx < e:
        # If anchor is inside moved block, skip.
        return logs, None

    block = [paras[i]._p for i in range(s, e)]
    anchor = paras[anchor_idx]._p
    parent = anchor.getparent()
    for el in block:
        if el.getparent() is parent:
            parent.remove(el)

    insert_pos = parent.index(anchor) + 1
    for el in block:
        parent.insert(insert_pos, el)
        insert_pos += 1

    logs.append(
        {
            "action": "move_prefix_toc_after_front_matter",
            "from": [s, e],
            "after_anchor": anchor_idx,
            "confidence": cls.confidence,
        }
    )
    moved_idx: list[int] = []
    paras_after = list(doc.paragraphs)
    for i, p in enumerate(paras_after):
        for el in block:
            if p._p is el:
                moved_idx.append(i)
                break
    moved_range: tuple[int, int] | None = None
    if moved_idx:
        moved_range = (min(moved_idx), max(moved_idx) + 1)
    return logs, moved_range


def ensure_marker_styles(doc: Document):
    for t in [
        T_EMPTY,
        T_BODY,
        T_TOC_TITLE,
        T_TOC_ENTRY,
        T_ABS_ZH_TITLE,
        T_ABS_ZH_BODY,
        T_KW_ZH,
        T_ABS_EN_TITLE,
        T_ABS_EN_BODY,
        T_KW_EN,
        T_H1,
        T_H2,
        T_H3,
        T_REF_TITLE,
        T_REF_ENTRY,
        T_ACK_TITLE,
        T_FIG_CAPTION,
        T_TABLE_CAPTION,
    ]:
        name = MARK_PREFIX + t
        try:
            doc.styles[name]
        except KeyError:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def write_type_markers(doc: Document, cls: Classification):
    ensure_marker_styles(doc)
    for p, tp in zip(doc.paragraphs, cls.types):
        p.style = doc.styles[MARK_PREFIX + tp]


def cleanup_marker_styles(doc: Document):
    styles_el = doc.styles.element
    remove_nodes = []
    for st in list(styles_el):
        if st.tag != qn("w:style"):
            continue
        style_id = st.get(qn("w:styleId")) or ""
        name_el = st.find(qn("w:name"))
        style_name = (name_el.get(qn("w:val")) if name_el is not None else "") or ""
        if style_id.startswith(MARK_PREFIX) or style_name.startswith(MARK_PREFIX):
            remove_nodes.append(st)
    for st in remove_nodes:
        styles_el.remove(st)


def apply_final_styles_from_markers(doc: Document, rules: dict) -> int:
    removed_numpr = 0
    bullet_prefix_re = re.compile(r"^\s*[\u25aa\u2022\u25cf\u25a0\u25c6\u25c7\u25e6\u00b7]+\s*")
    heading_align = {
        T_H1: rules.get("heading_1", {}).get("align", "center"),
        T_H2: rules.get("heading_2", {}).get("align", "left"),
        T_H3: rules.get("heading_3", {}).get("align", "left"),
    }

    for p in doc.paragraphs:
        if clear_paragraph_numbering(p):
            removed_numpr += 1
        p.paragraph_format.page_break_before = False

        style_name = p.style.name if p.style else ""
        text = (p.text or "").strip()
        in_table = is_in_table_cell(p)
        tp = style_name[len(MARK_PREFIX) :] if style_name.startswith(MARK_PREFIX) else None
        if tp is None:
            tp = T_EMPTY if not text else T_BODY

        # Remove artifact bullets generated by broken list metadata before styling.
        if tp not in {T_BODY, T_ABS_ZH_BODY, T_ABS_EN_BODY}:
            original = p.text or ""
            cleaned = bullet_prefix_re.sub("", original, count=1)
            if cleaned != original:
                p.text = cleaned
                text = cleaned.strip()

        # Collapse manual line-break noise in headings/captions/table cells.
        original2 = p.text or ""
        if in_table:
            cleaned2 = collapse_soft_breaks(original2, joiner="")
        elif tp in {
            T_ABS_ZH_TITLE,
            T_ABS_EN_TITLE,
            T_TOC_TITLE,
            T_TOC_ENTRY,
            T_H1,
            T_H2,
            T_H3,
            T_REF_TITLE,
            T_ACK_TITLE,
            T_KW_ZH,
            T_KW_EN,
            T_FIG_CAPTION,
            T_TABLE_CAPTION,
        }:
            cleaned2 = collapse_soft_breaks(original2, joiner="")
        else:
            cleaned2 = original2
        if cleaned2 != original2:
            p.text = cleaned2
            text = cleaned2.strip()

        if tp == T_EMPTY:
            p.style = doc.styles["Normal"]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Keep table content stable: no first-line indent, left align.
        if in_table:
            p.style = doc.styles["Normal"]
            tb = rules.get("table_body", {})
            set_run_style(p, tb.get("font", rules["body"]["font"]), tb.get("size_pt", rules["body"]["size_pt"]), False)
            set_para_format(
                p,
                align=tb.get("align", "left"),
                line_spacing=tb.get("line_spacing", rules["body"]["line_spacing"]),
                indent_chars=0,
                space_before=0,
                space_after=0,
            )
            continue

        if tp == T_ABS_ZH_TITLE:
            p.style = doc.styles["Title"]
            set_run_style(p, rules["abstract_title"]["font"], rules["abstract_title"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif tp == T_ABS_EN_TITLE:
            p.style = doc.styles["Title"]
            set_run_style(p, rules["english"]["font"], 18, True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif tp == T_TOC_TITLE:
            p.style = doc.styles["Title"]
            set_run_style(p, rules["toc"]["font"], rules["toc"]["title_size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=0, space_after=12)
            p.paragraph_format.page_break_before = True
        elif tp == T_REF_TITLE or tp == T_ACK_TITLE:
            p.style = doc.styles["Title"]
            set_run_style(p, rules["heading_1"]["font"], rules["heading_1"]["size_pt"], True)
            set_para_format(p, align="center", line_spacing=1.25, indent_chars=0, space_before=12, space_after=12)
            p.paragraph_format.page_break_before = True
        elif tp == T_H1:
            p.style = doc.styles["Heading 1"]
            set_run_style(p, rules["heading_1"]["font"], rules["heading_1"]["size_pt"], True)
            set_para_format(p, align=heading_align[T_H1], line_spacing=1.25, indent_chars=0, space_before=12, space_after=12)
        elif tp == T_H2:
            p.style = doc.styles["Heading 2"]
            set_run_style(p, rules["heading_2"]["font"], rules["heading_2"]["size_pt"], True)
            set_para_format(p, align=heading_align[T_H2], line_spacing=1.25, indent_chars=0, space_before=6, space_after=6)
        elif tp == T_H3:
            p.style = doc.styles["Heading 3"]
            set_run_style(p, rules["heading_3"]["font"], rules["heading_3"]["size_pt"], True)
            set_para_format(p, align=heading_align[T_H3], line_spacing=1.25, indent_chars=0, space_before=6, space_after=6)
        elif tp == T_FIG_CAPTION:
            cap = rules.get("figure_caption", {})
            p.style = doc.styles["Normal"]
            set_run_style(p, cap.get("font", rules["body"]["font"]), cap.get("size_pt", 10.5), bool(cap.get("bold", False)))
            set_para_format(p, align=cap.get("align", "center"), line_spacing=1.0, indent_chars=0, space_before=6, space_after=6)
        elif tp == T_TABLE_CAPTION:
            cap = rules.get("table_caption", {})
            p.style = doc.styles["Normal"]
            set_run_style(p, cap.get("font", rules["body"]["font"]), cap.get("size_pt", 10.5), bool(cap.get("bold", False)))
            set_para_format(p, align=cap.get("align", "center"), line_spacing=1.0, indent_chars=0, space_before=6, space_after=6)
        elif tp == T_TOC_ENTRY:
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["toc"]["font"], rules["toc"]["body_size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=0, space_after=0)
        elif tp == T_REF_ENTRY:
            p.style = doc.styles["Normal"]
            font = rules["body"]["font"]
            if re.search(r"[A-Za-z]{6,}", text) and not re.search(r"[\u4e00-\u9fff]", text):
                font = rules["english"]["font"]
            set_run_style(p, font, rules["body"]["size_pt"], False)
            set_para_format(
                p,
                align="left",
                line_spacing=rules["body"]["line_spacing"],
                indent_chars=0,
                space_before=0,
                space_after=0,
            )
        elif tp == T_KW_ZH:
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["body"]["font"], rules["body"]["size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=6, space_after=12)
        elif tp == T_KW_EN:
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["english"]["font"], rules["english"]["size_pt"], False)
            set_para_format(p, align="left", line_spacing=1.25, indent_chars=0, space_before=6, space_after=12)
        elif tp == T_ABS_EN_BODY:
            p.style = doc.styles["Normal"]
            set_run_style(p, rules["english"]["font"], rules["english"]["size_pt"], False)
            set_para_format(
                p,
                align="justify",
                line_spacing=rules["english"]["line_spacing"],
                indent_chars=2,
                space_before=0,
                space_after=0,
            )
        else:
            # BODY + ABS_ZH_BODY
            p.style = doc.styles["Normal"]
            font = rules["body"]["font"]
            if tp == T_ABS_EN_BODY or (re.search(r"[A-Za-z]{6,}", text) and not re.search(r"[\u4e00-\u9fff]", text)):
                font = rules["english"]["font"]
            set_run_style(p, font, rules["body"]["size_pt"], False)
            set_para_format(
                p,
                align=rules["body"].get("alignment", "justify"),
                line_spacing=rules["body"]["line_spacing"],
                indent_chars=rules["body"]["first_line_indent_chars"],
                space_before=0,
                space_after=0,
            )
    return removed_numpr


def _count_chars(doc: Document) -> int:
    texts = []
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)
    return len(re.sub(r"\s+", "", "\n".join(texts)))


@dataclass
class PipelineResult:
    output: str
    chars_no_space: int
    removed_numpr_count: int
    classification_confidence: float
    content_fingerprint_before: str = ""
    content_fingerprint_after: str = ""
    content_changed: bool = False
    logs: list[dict[str, Any]] = field(default_factory=list)


def run_pipeline(
    input_docx: str | Path,
    output_docx: str | Path,
    rules: dict,
    write_marker_dump: Path | None = None,
    enforce_content_guard: bool = True,
) -> PipelineResult:
    doc = Document(str(input_docx))
    fp_before = build_content_fingerprint(doc)
    base_result = setup_document_base(doc, rules)

    cls1 = classify_document(doc)
    logs = [
        {"action": "setup_document_base", **base_result},
        {"action": "classify_pass_1", "confidence": cls1.confidence, "notes": cls1.notes},
    ]
    reorder_logs, moved_toc_range = reorder_by_types(doc, cls1)
    logs.extend(reorder_logs)

    # Re-classify after possible move, then write markers.
    cls2 = classify_document(doc)
    if moved_toc_range is not None:
        s2, e2 = moved_toc_range
        paras2 = list(doc.paragraphs)
        for i in range(max(0, s2), min(len(cls2.types), e2)):
            t = (paras2[i].text or "").strip()
            if not t:
                continue
            cls2.types[i] = T_TOC_TITLE if is_toc_title(t) else T_TOC_ENTRY
        cls2.notes.append(f"force_toc_range={moved_toc_range}")
    logs.append({"action": "classify_pass_2", "confidence": cls2.confidence, "notes": cls2.notes})
    write_type_markers(doc, cls2)

    if write_marker_dump is not None:
        dump = []
        for i, p in enumerate(doc.paragraphs):
            t = (p.text or "").strip()
            tp = cls2.types[i]
            dump.append({"idx": i, "type": tp, "text": t[:120]})
        write_marker_dump.write_text(json.dumps(dump, ensure_ascii=False, indent=2), encoding="utf-8")

    removed_numpr_count = apply_final_styles_from_markers(doc, rules)
    logs.append({"action": "remove_numpr", "count": removed_numpr_count})
    cleanup_marker_styles(doc)
    logs.append({"action": "cleanup_marker_styles"})

    fp_after = build_content_fingerprint(doc)
    content_changed = fp_before != fp_after
    logs.append(
        {
            "action": "content_guard",
            "changed": content_changed,
            "fingerprint_before": fp_before,
            "fingerprint_after": fp_after,
        }
    )
    if content_changed and enforce_content_guard:
        raise ValueError("content guard failed: non-whitespace content changed")

    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return PipelineResult(
        output=str(output_docx),
        chars_no_space=_count_chars(doc),
        removed_numpr_count=removed_numpr_count,
        classification_confidence=cls2.confidence,
        content_fingerprint_before=fp_before,
        content_fingerprint_after=fp_after,
        content_changed=content_changed,
        logs=logs,
    )
