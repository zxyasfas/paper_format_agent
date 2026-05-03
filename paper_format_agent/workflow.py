from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Optional, TypedDict

from docx import Document

from .formatter import apply_formatting
from .llm_enhancer import config_from_env_or_args, enhance_docx_with_llm
from .postprocess import update_fields_with_libreoffice
from .rules import extract_rules_from_text
from .scorer import save_reports, score_document


class WorkflowState(TypedDict, total=False):
    args: dict[str, Any]
    format_text: str
    rules: dict[str, Any]
    author_bio_text: Optional[str]
    use_llm: bool
    auto_iterate: bool
    auto_iterate_allow_content_edit: bool
    target_score: float
    max_rounds: int
    strategy_plan: list[dict[str, Any]]
    strategy_index: int
    current_strategy_name: str
    round_no: int
    llm_cfg: Any
    llm_report: dict[str, Any]
    round_dir: str
    current_paper_file: str
    current_strict_content_fix: bool
    current_format_only: bool
    current_apply_llm_content_fixes: bool
    effective_paper_file: str
    output_docx: str
    format_result: dict[str, Any]
    libreoffice_fields_updated: bool
    current_score: float
    best_score: float
    best_round_no: int
    best_output_docx: str
    best_round_dir: str
    iteration_history: list[dict[str, Any]]
    continue_loop: bool
    stop_reason: str
    report: dict[str, Any]
    summary: dict[str, Any]


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)
    return "\n".join(texts)


def read_format_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    # For .doc/.pdf and other office formats, try LibreOffice conversion to docx.
    with tempfile.TemporaryDirectory() as td:
        out_dir = Path(td)
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=60,
            )
            converted = out_dir / (path.stem + ".docx")
            if converted.exists():
                return read_docx_text(converted)
        except Exception:
            pass
    return ""


def _load_author_bio_text(author_bio_file: str | None) -> str | None:
    if not author_bio_file:
        return None
    bio_path = Path(author_bio_file)
    if not bio_path.exists():
        raise FileNotFoundError(f"author bio file not found: {bio_path}")
    return read_format_text(bio_path)


def _build_llm_config(args: dict[str, Any]):
    return config_from_env_or_args(
        use_llm=bool(args.get("use_llm")),
        llm_api_key=args.get("llm_api_key"),
        llm_base_url=args.get("llm_base_url"),
        llm_model=args.get("llm_model"),
        llm_timeout=int(args.get("llm_timeout", 90)),
        apply_llm_content_fixes=bool(args.get("apply_llm_content_fixes", False)),
    )


def _save_json(path: Path, data: Any):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _default_strategy_plan(allow_content_edit: bool) -> list[dict[str, Any]]:
    plan = [
        {
            "name": "base_format_only",
            "strict_content_fix": False,
            "format_only": True,
            "apply_llm_content_fixes": False,
        },
        {
            "name": "strict_format_only",
            "strict_content_fix": True,
            "format_only": True,
            "apply_llm_content_fixes": False,
        },
    ]
    if allow_content_edit:
        plan.extend(
            [
                {
                    "name": "strict_content_rule_fix",
                    "strict_content_fix": True,
                    "format_only": False,
                    "apply_llm_content_fixes": False,
                },
                {
                    "name": "strict_content_rule_llm_fix",
                    "strict_content_fix": True,
                    "format_only": False,
                    "apply_llm_content_fixes": True,
                },
            ]
        )
    return plan


def _load_strategy_plan(config_path: str | None, allow_content_edit: bool) -> list[dict[str, Any]]:
    plan = _default_strategy_plan(allow_content_edit=allow_content_edit)
    if not config_path:
        return plan
    p = Path(config_path)
    if not p.exists():
        raise FileNotFoundError(f"strategy config not found: {p}")
    data = json.loads(p.read_text(encoding="utf-8"))
    raw_items = data.get("strategy_plan") if isinstance(data, dict) else data
    if not isinstance(raw_items, list):
        raise ValueError("strategy config must be a list or an object with 'strategy_plan' list.")

    parsed: list[dict[str, Any]] = []
    for i, item in enumerate(raw_items):
        if not isinstance(item, dict):
            continue
        parsed.append(
            {
                "name": str(item.get("name") or f"strategy_{i+1}"),
                "strict_content_fix": bool(item.get("strict_content_fix", False)),
                "format_only": bool(item.get("format_only", True)),
                "apply_llm_content_fixes": bool(item.get("apply_llm_content_fixes", False)),
            }
        )
    if not parsed:
        raise ValueError("strategy config has no valid strategy entries.")
    return parsed


def _promote_round_outputs(round_dir: Path, out_dir: Path):
    files = [
        "formatted_paper_95plus.docx",
        "modify_log.json",
        "format_report.json",
        "format_report.html",
        "llm_enhancement.json",
    ]
    for name in files:
        src = round_dir / name
        if src.exists():
            shutil.copy2(src, out_dir / name)


def _round_strategy_snapshot(state: WorkflowState) -> dict[str, Any]:
    return {
        "strict_content_fix": bool(state.get("current_strict_content_fix")),
        "format_only": bool(state.get("current_format_only", True)),
        "apply_llm_content_fixes": bool(state.get("current_apply_llm_content_fixes")),
    }


def run_legacy_pipeline(args: dict[str, Any]) -> dict[str, Any]:
    out_dir = Path(args["out_dir"])
    out_dir.mkdir(parents=True, exist_ok=True)

    format_text = read_format_text(args["format_file"])
    rules = extract_rules_from_text(format_text)
    _save_json(out_dir / "format_rules.json", rules)

    author_bio_text = _load_author_bio_text(args.get("author_bio_file"))

    llm_cfg = _build_llm_config(args)
    llm_report = enhance_docx_with_llm(
        paper_file=args["paper_file"],
        format_text=format_text,
        out_dir=out_dir,
        cfg=llm_cfg,
    )
    _save_json(out_dir / "llm_enhancement.json", llm_report)

    effective_paper_file = llm_report.get("output_docx") or args["paper_file"]
    output_docx = out_dir / "formatted_paper_95plus.docx"
    result = apply_formatting(
        effective_paper_file,
        output_docx,
        rules=rules,
        strict_content_fix=bool(args.get("strict_content_fix")),
        author_bio_text=author_bio_text,
        format_only=bool(args.get("format_only", True)),
    )
    if llm_report.get("actions"):
        for action in llm_report["actions"]:
            result["log"].append(
                {
                    "action": "llm_enhancement",
                    "target": action.get("target") or action.get("action", ""),
                    "before": None,
                    "after": json.dumps(action, ensure_ascii=False),
                    "note": f"model={llm_cfg.model}",
                }
            )

    updated = False if args.get("no_update_fields") else update_fields_with_libreoffice(output_docx)
    result["libreoffice_fields_updated"] = updated
    _save_json(out_dir / "modify_log.json", result["log"])

    report = score_document(output_docx, rules)
    report["libreoffice_fields_updated"] = updated
    report["llm_used"] = bool(llm_report.get("used"))
    report["llm_warnings"] = llm_report.get("warnings", [])
    save_reports(report, out_dir / "format_report.json", out_dir / "format_report.html")

    return {
        "output": str(output_docx),
        "score": report["score"],
        "chars_no_space": report["chars_no_space"],
        "libreoffice_fields_updated": updated,
        "llm_used": bool(llm_report.get("used")),
    }


def run_langgraph_pipeline(args: dict[str, Any]) -> dict[str, Any]:
    try:
        from langgraph.graph import END, StateGraph
    except Exception as e:
        raise RuntimeError(
            "LangGraph is not installed. Please run: pip install langgraph"
        ) from e

    out_dir = Path(args["out_dir"])
    out_dir.mkdir(parents=True, exist_ok=True)

    def node_prepare(state: WorkflowState) -> WorkflowState:
        format_text = read_format_text(args["format_file"])
        rules = extract_rules_from_text(format_text)
        _save_json(out_dir / "format_rules.json", rules)
        author_bio_text = _load_author_bio_text(args.get("author_bio_file"))
        use_llm = bool(args.get("use_llm"))
        auto_iterate = bool(args.get("auto_iterate", False))
        max_rounds = max(1, int(args.get("max_rounds", 3)))
        target_score = float(args.get("target_score", 95.0))
        allow_content = bool(args.get("auto_iterate_allow_content_edit", False))
        strategy_plan = _load_strategy_plan(args.get("strategy_config"), allow_content_edit=allow_content)
        user_initial = {
            "name": "user_initial",
            "strict_content_fix": bool(args.get("strict_content_fix", False)),
            "format_only": bool(args.get("format_only", True)),
            "apply_llm_content_fixes": bool(args.get("apply_llm_content_fixes", False)),
        }
        if not auto_iterate:
            strategy_plan = [user_initial]
        else:
            first = strategy_plan[0]
            if any(
                bool(first.get(k)) != bool(user_initial.get(k))
                for k in ("strict_content_fix", "format_only", "apply_llm_content_fixes")
            ):
                strategy_plan = [user_initial] + strategy_plan
        strategy_index = 0
        current_strategy = strategy_plan[strategy_index]
        round_no = 1
        return {
            "format_text": format_text,
            "rules": rules,
            "author_bio_text": author_bio_text,
            "use_llm": use_llm,
            "auto_iterate": auto_iterate,
            "auto_iterate_allow_content_edit": allow_content,
            "target_score": target_score,
            "max_rounds": max_rounds,
            "strategy_plan": strategy_plan,
            "strategy_index": strategy_index,
            "current_strategy_name": current_strategy["name"],
            "round_no": round_no,
            "current_paper_file": str(args["paper_file"]),
            "current_strict_content_fix": bool(current_strategy["strict_content_fix"]),
            "current_format_only": bool(current_strategy["format_only"]),
            "current_apply_llm_content_fixes": bool(current_strategy["apply_llm_content_fixes"]),
            "best_score": -1.0,
            "best_round_no": 0,
            "iteration_history": [],
            "continue_loop": False,
            "stop_reason": "",
        }

    def node_llm(state: WorkflowState) -> WorkflowState:
        round_no = int(state.get("round_no", 1))
        round_dir = out_dir / f"round_{round_no}"
        round_dir.mkdir(parents=True, exist_ok=True)
        llm_cfg = config_from_env_or_args(
            use_llm=bool(state.get("use_llm")),
            llm_api_key=args.get("llm_api_key"),
            llm_base_url=args.get("llm_base_url"),
            llm_model=args.get("llm_model"),
            llm_timeout=int(args.get("llm_timeout", 90)),
            apply_llm_content_fixes=bool(state.get("current_apply_llm_content_fixes", False)),
        )
        llm_report = enhance_docx_with_llm(
            paper_file=state.get("current_paper_file", str(args["paper_file"])),
            format_text=state.get("format_text", ""),
            out_dir=round_dir,
            cfg=llm_cfg,
        )
        _save_json(round_dir / "llm_enhancement.json", llm_report)
        effective_paper_file = llm_report.get("output_docx") or state.get("current_paper_file", str(args["paper_file"]))
        return {
            "round_dir": str(round_dir),
            "llm_cfg": llm_cfg,
            "llm_report": llm_report,
            "effective_paper_file": str(effective_paper_file),
        }

    def node_format(state: WorkflowState) -> WorkflowState:
        round_dir = Path(state["round_dir"])
        output_docx = round_dir / "formatted_paper_95plus.docx"
        result = apply_formatting(
            state.get("effective_paper_file", state.get("current_paper_file", str(args["paper_file"]))),
            output_docx,
            rules=state["rules"],
            strict_content_fix=bool(state.get("current_strict_content_fix")),
            author_bio_text=state.get("author_bio_text"),
            format_only=bool(state.get("current_format_only", True)),
        )
        llm_report = state.get("llm_report", {})
        llm_cfg = state.get("llm_cfg")
        if llm_report.get("actions"):
            for action in llm_report["actions"]:
                result["log"].append(
                    {
                        "action": "llm_enhancement",
                        "target": action.get("target") or action.get("action", ""),
                        "before": None,
                        "after": json.dumps(action, ensure_ascii=False),
                        "note": f"model={getattr(llm_cfg, 'model', '')}",
                    }
                )
        updated = False if args.get("no_update_fields") else update_fields_with_libreoffice(output_docx)
        result["libreoffice_fields_updated"] = updated
        _save_json(round_dir / "modify_log.json", result["log"])
        return {
            "output_docx": str(output_docx),
            "format_result": result,
            "libreoffice_fields_updated": updated,
        }

    def node_score(state: WorkflowState) -> WorkflowState:
        round_no = int(state.get("round_no", 1))
        round_dir = Path(state["round_dir"])
        output_docx = Path(state["output_docx"])
        report = score_document(output_docx, state["rules"])
        report["libreoffice_fields_updated"] = bool(state.get("libreoffice_fields_updated"))
        llm_report = state.get("llm_report", {})
        report["llm_used"] = bool(llm_report.get("used"))
        report["llm_warnings"] = llm_report.get("warnings", [])
        save_reports(report, round_dir / "format_report.json", round_dir / "format_report.html")
        current_score = float(report["score"])
        history = list(state.get("iteration_history", []))
        history.append(
            {
                "round": round_no,
                "score": current_score,
                "chars_no_space": report["chars_no_space"],
                "strategy_name": state.get("current_strategy_name"),
                "strategy": _round_strategy_snapshot(state),
                "llm_used": bool(llm_report.get("used")),
                "llm_warnings": llm_report.get("warnings", []),
                "round_dir": str(round_dir),
            }
        )
        best_score = float(state.get("best_score", -1.0))
        best_round_no = int(state.get("best_round_no", 0))
        best_output_docx = state.get("best_output_docx", "")
        best_round_dir = state.get("best_round_dir", "")
        if current_score >= best_score:
            best_score = current_score
            best_round_no = round_no
            best_output_docx = str(output_docx)
            best_round_dir = str(round_dir)
            _promote_round_outputs(round_dir, out_dir)

        need_next = (
            bool(state.get("auto_iterate"))
            and current_score < float(state.get("target_score", 95.0))
            and round_no < int(state.get("max_rounds", 1))
        )
        stop_reason = state.get("stop_reason", "")
        if not need_next:
            if current_score >= float(state.get("target_score", 95.0)):
                stop_reason = "target_score_reached"
            elif not bool(state.get("auto_iterate")):
                stop_reason = "auto_iterate_disabled"
            else:
                stop_reason = "max_rounds_reached"

        summary = {
            "output": str(output_docx),
            "score": current_score,
            "chars_no_space": report["chars_no_space"],
            "libreoffice_fields_updated": bool(state.get("libreoffice_fields_updated")),
            "llm_used": bool(llm_report.get("used")),
            "round": round_no,
        }
        return {
            "report": report,
            "summary": summary,
            "current_score": current_score,
            "iteration_history": history,
            "best_score": best_score,
            "best_round_no": best_round_no,
            "best_output_docx": best_output_docx,
            "best_round_dir": best_round_dir,
            "continue_loop": need_next,
            "stop_reason": stop_reason,
        }

    def node_tune(state: WorkflowState) -> WorkflowState:
        if not bool(state.get("continue_loop")):
            return {"continue_loop": False}

        round_no = int(state.get("round_no", 1))
        next_round = round_no + 1
        strategy_plan = list(state.get("strategy_plan", []))
        strategy_index = int(state.get("strategy_index", 0))
        next_index = strategy_index + 1
        if next_index >= len(strategy_plan):
            return {
                "continue_loop": False,
                "stop_reason": "no_more_strategies",
            }

        next_strategy = strategy_plan[next_index]
        strategy_note = str(next_strategy.get("name", f"strategy_{next_index+1}"))

        next_input = state.get("best_output_docx") or state.get("output_docx") or state.get("current_paper_file")
        if not next_input:
            return {
                "continue_loop": False,
                "stop_reason": "missing_next_input",
            }

        history = list(state.get("iteration_history", []))
        if history:
            history[-1]["next_strategy"] = strategy_note

        return {
            "round_no": next_round,
            "strategy_index": next_index,
            "current_strategy_name": strategy_note,
            "current_paper_file": str(next_input),
            "current_strict_content_fix": bool(next_strategy.get("strict_content_fix", False)),
            "current_format_only": bool(next_strategy.get("format_only", True)),
            "current_apply_llm_content_fixes": bool(next_strategy.get("apply_llm_content_fixes", False)),
            "continue_loop": True,
            "iteration_history": history,
        }

    def node_finalize(state: WorkflowState) -> WorkflowState:
        best_output_docx = state.get("best_output_docx") or state.get("output_docx")
        best_score = float(state.get("best_score", state.get("current_score", 0.0)))
        best_round_no = int(state.get("best_round_no", state.get("round_no", 1)))
        stop_reason = state.get("stop_reason", "completed")
        history = list(state.get("iteration_history", []))
        best_record = next((h for h in history if int(h.get("round", 0)) == best_round_no), None)
        iteration_report = {
            "target_score": float(state.get("target_score", 95.0)),
            "max_rounds": int(state.get("max_rounds", 1)),
            "auto_iterate": bool(state.get("auto_iterate")),
            "auto_iterate_allow_content_edit": bool(state.get("auto_iterate_allow_content_edit")),
            "strategy_plan": state.get("strategy_plan", []),
            "stop_reason": stop_reason,
            "best_score": best_score,
            "best_round_no": best_round_no,
            "best_output_docx": best_output_docx,
            "history": history,
        }
        _save_json(out_dir / "iteration_report.json", iteration_report)
        summary = {
            "output": str(out_dir / "formatted_paper_95plus.docx") if (out_dir / "formatted_paper_95plus.docx").exists() else str(best_output_docx),
            "score": best_score,
            "best_round": best_round_no,
            "iterations": len(history),
            "target_score": float(state.get("target_score", 95.0)),
            "stop_reason": stop_reason,
            "chars_no_space": (best_record or {}).get("chars_no_space"),
            "libreoffice_fields_updated": state.get("summary", {}).get("libreoffice_fields_updated"),
            "llm_used": (best_record or {}).get("llm_used"),
        }
        return {"summary": summary}

    def route_after_tune(state: WorkflowState) -> str:
        return "loop" if bool(state.get("continue_loop")) else "finalize"

    graph = StateGraph(WorkflowState)
    graph.add_node("prepare", node_prepare)
    graph.add_node("llm", node_llm)
    graph.add_node("format", node_format)
    graph.add_node("score", node_score)
    graph.add_node("tune", node_tune)
    graph.add_node("finalize", node_finalize)
    graph.set_entry_point("prepare")
    graph.add_edge("prepare", "llm")
    graph.add_edge("llm", "format")
    graph.add_edge("format", "score")
    graph.add_edge("score", "tune")
    graph.add_conditional_edges("tune", route_after_tune, {"loop": "llm", "finalize": "finalize"})
    graph.add_edge("finalize", END)
    app = graph.compile()

    result_state = app.invoke({"args": args})
    return result_state["summary"]
