from __future__ import annotations

import json
import re
from html import escape
from pathlib import Path
from typing import Any

from docx import Document

from .pipeline import (
    MARK_PREFIX,
    is_abstract_title,
    is_english_abstract_title,
    is_keyword_en,
    is_keyword_zh,
    looks_like_toc_entry,
    is_toc_title,
    normalize_text,
)


DIAGNOSTIC_CATALOG: dict[str, dict[str, str]] = {
    "missing_zh_abs": {
        "category": "required_sections",
        "severity": "high",
        "summary": "Chinese abstract section is missing.",
        "suggested_fix": "Add a standalone Chinese abstract title before body text.",
    },
    "missing_zh_keywords": {
        "category": "required_sections",
        "severity": "medium",
        "summary": "Chinese keywords line is missing.",
        "suggested_fix": "Add a keywords line after the Chinese abstract, for example 'Keywords: ...' in the required language.",
    },
    "missing_en_abs": {
        "category": "required_sections",
        "severity": "medium",
        "summary": "English abstract section is missing.",
        "suggested_fix": "Add an 'Abstract' section after the Chinese abstract and keywords when the template requires it.",
    },
    "missing_en_keywords": {
        "category": "required_sections",
        "severity": "medium",
        "summary": "English keywords line is missing.",
        "suggested_fix": "Add a 'Keywords:' line after the English abstract.",
    },
    "missing_toc_title": {
        "category": "required_sections",
        "severity": "low",
        "summary": "Table of contents title is missing.",
        "suggested_fix": "Insert a table of contents title using the wording required by the format guide.",
    },
    "abstract_after_body": {
        "category": "front_matter_order",
        "severity": "high",
        "summary": "Abstract appears after body content.",
        "suggested_fix": "Move the abstract and keyword sections before the introduction or first body heading.",
    },
    "prefix_manual_toc_before_abs": {
        "category": "front_matter_order",
        "severity": "high",
        "summary": "Manual table of contents appears before the abstract.",
        "suggested_fix": "Move the table of contents after the abstract and keyword block, or regenerate the document with TOC reordering enabled.",
    },
    "en_abs_before_zh_abs": {
        "category": "front_matter_order",
        "severity": "medium",
        "summary": "English abstract appears before Chinese abstract.",
        "suggested_fix": "Place the Chinese abstract first, followed by Chinese keywords, English abstract, and English keywords.",
    },
    "kw_zh_before_abs": {
        "category": "front_matter_order",
        "severity": "medium",
        "summary": "Chinese keywords appear before the Chinese abstract.",
        "suggested_fix": "Move Chinese keywords directly after the Chinese abstract body.",
    },
    "kw_en_before_en_abs": {
        "category": "front_matter_order",
        "severity": "medium",
        "summary": "English keywords appear before the English abstract.",
        "suggested_fix": "Move English keywords directly after the English abstract body.",
    },
    "toc_after_body": {
        "category": "front_matter_order",
        "severity": "medium",
        "summary": "Table of contents appears after body content.",
        "suggested_fix": "Move the table of contents into the front matter before the first body chapter.",
    },
    "numpr_left": {
        "category": "docx_numbering",
        "severity": "medium",
        "summary": "Word numbering metadata remains in paragraphs.",
        "suggested_fix": "Clear list numbering from affected paragraphs and reapply deterministic paragraph styles.",
    },
    "marker_left": {
        "category": "internal_cleanup",
        "severity": "medium",
        "summary": "Internal marker styles remain in the formatted document.",
        "suggested_fix": "Run marker cleanup before returning the formatted DOCX.",
    },
    "heading_body_leak": {
        "category": "style_leak",
        "severity": "high",
        "summary": "Long body-like paragraphs are styled as headings.",
        "suggested_fix": "Convert long heading-styled paragraphs back to body style and keep heading styles for short section titles only.",
    },
    "toc_heading_leak": {
        "category": "style_leak",
        "severity": "medium",
        "summary": "TOC entries are still styled as headings.",
        "suggested_fix": "Apply normal TOC-entry formatting to TOC lines instead of Heading styles.",
    },
    "blank_page_risk": {
        "category": "layout",
        "severity": "medium",
        "summary": "Forced page breaks may create sparse or blank pages.",
        "suggested_fix": "Review page-break-before settings near front-matter and heading paragraphs.",
    },
    "content_loss_vs_baseline": {
        "category": "content_preservation",
        "severity": "high",
        "summary": "Formatted output appears to contain less content than the baseline.",
        "suggested_fix": "Compare the input and output documents before accepting the format run.",
    },
    "char_below_min": {
        "category": "template_rules",
        "severity": "medium",
        "summary": "Document length is below the minimum required by the format guide.",
        "suggested_fix": "Verify whether the format guide's minimum character count applies to this manuscript.",
    },
}


def _unknown_diagnostic_severity(value: int | float) -> str:
    if value >= 15:
        return "high"
    if value >= 8:
        return "medium"
    return "low"


def build_diagnostics(
    penalties: list[dict[str, Any]], features: dict[str, Any] | None = None
) -> list[dict[str, Any]]:
    diagnostics: list[dict[str, Any]] = []
    features = features or {}
    for penalty in penalties:
        name = str(penalty.get("name", "unknown"))
        value = penalty.get("value", 0)
        catalog = DIAGNOSTIC_CATALOG.get(name, {})
        diagnostics.append(
            {
                "name": name,
                "category": catalog.get("category", "uncategorized"),
                "severity": catalog.get("severity", _unknown_diagnostic_severity(float(value or 0))),
                "penalty": value,
                "summary": catalog.get("summary", f"Formatting check failed: {name}."),
                "suggested_fix": catalog.get(
                    "suggested_fix",
                    "Review this check in the format guide and adjust the document or rule extractor.",
                ),
                "evidence": _diagnostic_evidence(name, features),
            }
        )
    return diagnostics


def _diagnostic_evidence(name: str, features: dict[str, Any]) -> dict[str, Any]:
    keys_by_name = {
        "abstract_after_body": ["idx_abs", "idx_intro"],
        "toc_after_body": ["idx_toc", "idx_intro"],
        "prefix_manual_toc_before_abs": [
            "prefix_non_empty_before_abs",
            "prefix_toc_like_before_abs",
            "prefix_toc_like_ratio_before_abs",
        ],
        "numpr_left": ["numpr_left"],
        "marker_left": ["marker_left"],
        "heading_body_leak": ["heading_body_leak"],
        "toc_heading_leak": ["toc_heading_leak"],
        "blank_page_risk": ["blank_page_risk"],
        "content_loss_vs_baseline": ["baseline_chars_no_space", "chars_ratio_vs_baseline"],
        "char_below_min": ["chars_no_space", "min_chars"],
    }
    return {key: features.get(key) for key in keys_by_name.get(name, [])}


def _count_chars(doc: Document) -> int:
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    parts.append(c.text)
    return len(re.sub(r"\s+", "", "\n".join(parts)))


def _first_index(doc: Document, pred) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if pred((p.text or "").strip()):
            return i
    return None


def _is_intro_or_h1(text: str) -> bool:
    s = (text or "").strip()
    n = normalize_text(s)
    return ("绪论" in n) or ("引言" in n) or bool(re.match(r"^第[一二三四五六七八九十0-9]+章", s)) or bool(re.match(r"^[一二三四五六七八九十]+、", s))


def _page_break_blank_risk(doc: Document) -> int:
    """
    Approximate visual blank-page risk by checking forced page breaks
    followed by too-little textual content.
    """
    risk = 0
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if not p.paragraph_format.page_break_before:
            continue
        collected = 0
        seen = 0
        for j in range(i + 1, len(paras)):
            t = (paras[j].text or "").strip()
            if paras[j].paragraph_format.page_break_before and seen > 0:
                break
            if not t:
                continue
            seen += 1
            collected += len(normalize_text(t))
            if seen >= 4:
                break
        if seen > 0 and collected < 50:
            risk += 1
    return risk


def _prefix_toc_stats(doc: Document, idx_abs: int | None) -> dict[str, float]:
    if idx_abs is None or idx_abs <= 0:
        return {"non_empty": 0, "toc_like": 0, "long_body_like": 0, "toc_like_ratio": 0.0, "is_manual_toc": 0}
    non_empty = 0
    toc_like = 0
    long_body_like = 0
    for p in doc.paragraphs[:idx_abs]:
        t = (p.text or "").strip()
        if not t:
            continue
        non_empty += 1
        nt = normalize_text(t)
        if len(nt) >= 80 or (len(nt) >= 56 and len(re.findall(r"[。！？!?；;：:]", nt)) >= 2):
            long_body_like += 1
        if looks_like_toc_entry(t):
            toc_like += 1
    ratio = (toc_like / non_empty) if non_empty else 0.0
    is_manual_toc = int(non_empty >= 8 and ratio >= 0.72 and long_body_like <= max(1, non_empty // 12))
    return {
        "non_empty": non_empty,
        "toc_like": toc_like,
        "long_body_like": long_body_like,
        "toc_like_ratio": ratio,
        "is_manual_toc": is_manual_toc,
    }


def _toc_heading_leak(doc: Document, idx_toc: int | None) -> int:
    """Count heading-styled TOC lines after TOC title."""
    if idx_toc is None:
        return 0

    paras = list(doc.paragraphs)
    heading_like = 0
    short_non_empty = 0
    empty_run = 0

    for j in range(idx_toc + 1, len(paras)):
        p = paras[j]
        t = (p.text or "").strip()
        if not t:
            empty_run += 1
            if short_non_empty >= 5 and empty_run >= 3:
                break
            continue

        empty_run = 0
        nt = normalize_text(t)
        if len(nt) <= 80:
            short_non_empty += 1
            style_name = (p.style.name if p.style else "")
            if style_name.startswith("Heading"):
                heading_like += 1
        else:
            # Long body paragraph likely means TOC area ended.
            if short_non_empty >= 5:
                break

        if short_non_empty >= 80:
            break

    if short_non_empty < 5:
        return 0
    if heading_like < 3:
        return 0
    return heading_like


def score_document(
    docx_path: str | Path,
    rules: dict,
    calibration_file: str | Path | None = None,
    baseline_docx: str | Path | None = None,
    enforce_required_sections: bool = False,
) -> dict[str, Any]:
    doc = Document(str(docx_path))
    chars = _count_chars(doc)
    required_cfg = dict(rules.get("required_sections", {}))
    cfg_req_zh_abs = bool(required_cfg.get("zh_abstract", True))
    cfg_req_zh_kw = bool(required_cfg.get("zh_keywords", True))
    cfg_req_en_abs = bool(required_cfg.get("en_abstract", False))
    cfg_req_en_kw = bool(required_cfg.get("en_keywords", False))
    cfg_req_toc = bool(required_cfg.get("toc", False))

    idx_abs = _first_index(doc, is_abstract_title)
    idx_en_abs = _first_index(doc, is_english_abstract_title)
    idx_kw_zh = _first_index(doc, is_keyword_zh)
    idx_kw_en = _first_index(doc, is_keyword_en)
    idx_toc = _first_index(doc, is_toc_title)
    idx_intro = _first_index(doc, _is_intro_or_h1)

    baseline_presence = {"zh_abs": False, "zh_kw": False, "en_abs": False, "en_kw": False, "toc": False}
    baseline_chars = None
    if baseline_docx:
        p = Path(baseline_docx)
        if p.exists():
            base_doc = Document(str(p))
            baseline_presence = {
                "zh_abs": _first_index(base_doc, is_abstract_title) is not None,
                "zh_kw": _first_index(base_doc, is_keyword_zh) is not None,
                "en_abs": _first_index(base_doc, is_english_abstract_title) is not None,
                "en_kw": _first_index(base_doc, is_keyword_en) is not None,
                "toc": _first_index(base_doc, is_toc_title) is not None,
            }
            baseline_chars = _count_chars(base_doc)

    if enforce_required_sections:
        req_zh_abs = cfg_req_zh_abs
        req_zh_kw = cfg_req_zh_kw
        req_en_abs = cfg_req_en_abs
        req_en_kw = cfg_req_en_kw
        req_toc = cfg_req_toc
    else:
        req_zh_abs = baseline_presence["zh_abs"]
        req_zh_kw = baseline_presence["zh_kw"]
        req_en_abs = baseline_presence["en_abs"]
        req_en_kw = baseline_presence["en_kw"]
        req_toc = baseline_presence["toc"]

    numpr_left = sum(1 for p in doc.paragraphs if "w:numPr" in p._p.xml)
    marker_left = sum(1 for p in doc.paragraphs if (p.style and p.style.name.startswith(MARK_PREFIX)))
    heading_body_leak = sum(
        1
        for p in doc.paragraphs
        if (p.style and p.style.name.startswith("Heading")) and len(normalize_text((p.text or "").strip())) > 70
    )
    blank_risk = _page_break_blank_risk(doc)
    prefix_toc = _prefix_toc_stats(doc, idx_abs)
    toc_heading_leak = _toc_heading_leak(doc, idx_toc)

    penalties: list[dict[str, Any]] = []

    if req_zh_abs and idx_abs is None:
        penalties.append({"name": "missing_zh_abs", "value": 25})
    if req_zh_kw and idx_kw_zh is None:
        penalties.append({"name": "missing_zh_keywords", "value": 10})
    if req_en_abs and idx_en_abs is None:
        penalties.append({"name": "missing_en_abs", "value": 12})
    if req_en_kw and idx_kw_en is None:
        penalties.append({"name": "missing_en_keywords", "value": 10})
    if req_toc and idx_toc is None:
        penalties.append({"name": "missing_toc_title", "value": 6})

    if idx_abs is not None and idx_intro is not None and idx_abs > idx_intro:
        penalties.append({"name": "abstract_after_body", "value": 20})
    if prefix_toc["is_manual_toc"]:
        penalties.append({"name": "prefix_manual_toc_before_abs", "value": 22})
    if idx_en_abs is not None and idx_abs is not None and idx_en_abs < idx_abs:
        penalties.append({"name": "en_abs_before_zh_abs", "value": 8})
    if idx_kw_zh is not None and idx_abs is not None and idx_kw_zh < idx_abs:
        penalties.append({"name": "kw_zh_before_abs", "value": 6})
    if idx_kw_en is not None and idx_en_abs is not None and idx_kw_en < idx_en_abs:
        penalties.append({"name": "kw_en_before_en_abs", "value": 6})
    if idx_toc is not None and idx_intro is not None and idx_toc > idx_intro:
        penalties.append({"name": "toc_after_body", "value": 8})

    if numpr_left > 0:
        penalties.append({"name": "numpr_left", "value": min(20, numpr_left // 2 + 4)})
    if marker_left > 0:
        penalties.append({"name": "marker_left", "value": min(10, marker_left)})
    if heading_body_leak > 0:
        penalties.append({"name": "heading_body_leak", "value": min(16, heading_body_leak * 2)})
    if toc_heading_leak > 0:
        penalties.append({"name": "toc_heading_leak", "value": min(24, toc_heading_leak * 2)})
    if blank_risk > 0:
        penalties.append({"name": "blank_page_risk", "value": min(15, blank_risk * 3)})

    # Default mode checks content preservation, not absolute length thresholds.
    if baseline_chars is not None and baseline_chars > 0:
        char_ratio = chars / baseline_chars
        if char_ratio < 0.95:
            loss_ratio = max(0.0, 1.0 - char_ratio)
            penalties.append({"name": "content_loss_vs_baseline", "value": min(15, int(loss_ratio * 50))})
    else:
        char_ratio = None

    min_chars = int(rules.get("min_total_chars_no_space", 10000))
    if enforce_required_sections and chars < min_chars:
        penalties.append({"name": "char_below_min", "value": 15})

    base_score = 100.0
    penalty_sum = float(sum(x["value"] for x in penalties))
    quality_score = max(0.0, min(100.0, base_score - penalty_sum))

    calibrated_score = quality_score
    calibration = None
    if calibration_file:
        p = Path(calibration_file)
        if p.exists():
            try:
                calibration = json.loads(p.read_text(encoding="utf-8"))
                scale = float(calibration.get("scale", 1.0))
                offset = float(calibration.get("offset", 0.0))
                calibrated_score = max(0.0, min(100.0, quality_score * scale + offset))
            except Exception:
                calibration = None

    features = {
        "chars_no_space": chars,
        "idx_abs": idx_abs,
        "idx_en_abs": idx_en_abs,
        "idx_kw_zh": idx_kw_zh,
        "idx_kw_en": idx_kw_en,
        "idx_toc": idx_toc,
        "idx_intro": idx_intro,
        "enforce_required_sections": enforce_required_sections,
        "baseline_has_zh_abs": baseline_presence["zh_abs"],
        "baseline_has_zh_keywords": baseline_presence["zh_kw"],
        "baseline_has_en_abs": baseline_presence["en_abs"],
        "baseline_has_en_keywords": baseline_presence["en_kw"],
        "baseline_has_toc": baseline_presence["toc"],
        "required_zh_abs": req_zh_abs,
        "required_zh_keywords": req_zh_kw,
        "required_en_abs": req_en_abs,
        "required_en_keywords": req_en_kw,
        "required_toc": req_toc,
        "baseline_chars_no_space": baseline_chars,
        "chars_ratio_vs_baseline": (round(float(char_ratio), 4) if char_ratio is not None else None),
        "numpr_left": numpr_left,
        "marker_left": marker_left,
        "heading_body_leak": heading_body_leak,
        "toc_heading_leak": toc_heading_leak,
        "blank_page_risk": blank_risk,
        "prefix_non_empty_before_abs": prefix_toc["non_empty"],
        "prefix_toc_like_before_abs": prefix_toc["toc_like"],
        "prefix_long_body_like_before_abs": prefix_toc["long_body_like"],
        "prefix_toc_like_ratio_before_abs": round(float(prefix_toc["toc_like_ratio"]), 3),
        "prefix_manual_toc_before_abs": bool(prefix_toc["is_manual_toc"]),
        "min_chars": min_chars,
    }

    return {
        "score": round(calibrated_score, 1),
        "raw_quality_score": round(quality_score, 1),
        "chars_no_space": chars,
        "features": features,
        "penalties": penalties,
        "diagnostics": build_diagnostics(penalties, features),
        "calibration": calibration,
    }


def save_reports(report: dict[str, Any], out_json: str | Path, out_html: str | Path):
    out_json = Path(out_json)
    out_html = Path(out_html)
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    rows = []
    for p in report.get("penalties", []):
        rows.append(
            f"<tr><td>{escape(str(p['name']))}</td><td>-{escape(str(p['value']))}</td></tr>"
        )
    diagnostics = report.get("diagnostics")
    if diagnostics is None:
        diagnostics = build_diagnostics(report.get("penalties", []), report.get("features", {}))
    diag_rows = []
    for item in diagnostics:
        diag_rows.append(
            "<tr>"
            f"<td>{escape(str(item.get('severity', '')))}</td>"
            f"<td>{escape(str(item.get('name', '')))}</td>"
            f"<td>{escape(str(item.get('summary', '')))}</td>"
            f"<td>{escape(str(item.get('suggested_fix', '')))}</td>"
            "</tr>"
        )
    feat_rows = []
    for k, v in report.get("features", {}).items():
        feat_rows.append(f"<tr><td>{escape(str(k))}</td><td>{escape(str(v))}</td></tr>")

    # 检查是否有排版前后对比数据
    has_comparison = "score_before" in report and "score_after" in report
    
    if has_comparison:
        score_before = report['score_before']
        score_after = report['score_after']
        improvement = report.get('score_improvement', score_after - score_before)
        chars_before = report.get('chars_no_space_before', report.get('chars_no_space', 0))
        chars_after = report.get('chars_no_space_after', report.get('chars_no_space', 0))
        improvement_color = "#4CAF50" if improvement >= 0 else "#F44336"
        improvement_sign = "+" if improvement > 0 else ""
        
        comparison_html = f"""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 20px; border-radius: 10px; color: white; margin: 20px 0;">
            <h2 style="margin-top: 0; color: white;">📊 排版效果对比</h2>
            <div style="display: flex; justify-content: space-around; text-align: center;">
                <div style="padding: 15px;">
                    <div style="font-size: 14px; opacity: 0.9;">排版前</div>
                    <div style="font-size: 36px; font-weight: bold;">{score_before}</div>
                    <div style="font-size: 12px; opacity: 0.8;">{chars_before:,} 字符</div>
                </div>
                <div style="padding: 15px; display: flex; align-items: center;">
                    <div style="font-size: 48px;">→</div>
                </div>
                <div style="padding: 15px;">
                    <div style="font-size: 14px; opacity: 0.9;">排版后</div>
                    <div style="font-size: 36px; font-weight: bold;">{score_after}</div>
                    <div style="font-size: 12px; opacity: 0.8;">{chars_after:,} 字符</div>
                </div>
                <div style="padding: 15px;">
                    <div style="font-size: 14px; opacity: 0.9;">提升</div>
                    <div style="font-size: 36px; font-weight: bold; color: {improvement_color};">{improvement_sign}{improvement:.1f}</div>
                    <div style="font-size: 12px; opacity: 0.8;">分</div>
                </div>
            </div>
        </div>
        """
        score_display = f"{score_after}"
    else:
        comparison_html = ""
        score_display = f"{report['score']}"

    html = (
        "<!doctype html><html><head><meta charset='utf-8'><title>V3质量评分</title>"
        "<style>"
        "body{font-family:Arial,'Microsoft YaHei';max-width:1000px;margin:24px auto;line-height:1.6;background:#f5f5f5}"
        ".container{background:white;padding:30px;border-radius:10px;box-shadow:0 2px 10px rgba(0,0,0,0.1)}"
        "table{border-collapse:collapse;width:100%;margin:15px 0}"
        "td,th{border:1px solid #ddd;padding:10px;text-align:left}"
        "th{background:#f5f5f5;font-weight:bold}"
        "tr:hover{background:#f9f9f9}"
        "h1{color:#333;border-bottom:3px solid #667eea;padding-bottom:15px}"
        "h2{color:#555;margin-top:30px}"
        ".score-box{display:inline-block;padding:15px 30px;background:#667eea;color:white;border-radius:8px;font-size:24px;font-weight:bold}"
        "</style></head><body>"
        "<div class='container'>"
        "<h1>📄 Format Quality Score Report</h1><p><em>Review the penalties and diagnostics below to understand what failed and what to fix next.</em></p>"
        + comparison_html +
        f"<p><b>Final Score:</b> <span class='score-box'>{score_display}</span> / 100</p>"
        f"<p><b>Raw Quality Score:</b> {report.get('raw_quality_score', report['score'])}</p>"
        "<h2>⚠️ Penalties — What Failed</h2><table><tr><th>Check</th><th>Points Deducted</th></tr>"
        + "".join(rows)
        + "</table>"
        + "<h2>🔧 Diagnostics — What to Fix Next</h2><table><tr><th>Severity</th><th>Check</th><th>Problem</th><th>Suggested Fix</th></tr>"
        + "".join(diag_rows)
        + "</table>"
        + "<h2>📋 Feature Details</h2><table><tr><th>Feature</th><th>Value</th></tr>"
        + "".join(feat_rows)
        + "</table></div></body></html>"
    )
    out_html.write_text(html, encoding="utf-8")
